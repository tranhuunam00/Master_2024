import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def add_image_with_caption(doc, img_name, caption_text, width_inches):
    img_path = os.path.join(r"c:\Users\ADMIN\OneDrive\Máy tính\Master_2024\Cpap\document", img_name)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(img_path, width=Inches(width_inches))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(8)
        run_cap = p_cap.add_run(caption_text)
        run_cap.italic = True
        run_cap.font.size = Pt(10)
    else:
        print(f"Warning: Image {img_name} not found at {img_path}")

def create_document():
    doc = Document()
    
    # 1. Page Margins Setup (Left 3cm, Right 2cm, Top 2cm, Bottom 2cm)
    # 1 inch = 2.54 cm -> 3cm = 1.18 in, 2cm = 0.79 in
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.79)
        
    # 2. Typography and Styles Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    # Custom Paragraph format for Normal style
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # 3. Header Administrative Block
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header_format = p_header.paragraph_format
    p_header_format.space_after = Pt(2)
    
    run_nation1 = p_header.add_run("SOCIALIST REPUBLIC OF VIETNAM\n")
    run_nation1.bold = True
    run_nation1.font.size = Pt(12)
    
    run_nation2 = p_header.add_run("Independence - Freedom - Happiness\n")
    run_nation2.bold = True
    run_nation2.font.size = Pt(13)
    
    run_line = p_header.add_run("---------------o0o---------------")
    run_line.font.size = Pt(10)
    
    # 4. Title of the Patent
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(18)
    p_title.paragraph_format.space_after = Pt(18)
    
    run_title = p_title.add_run("DESCRIPTION OF UTILITY SOLUTION\n")
    run_title.bold = True
    run_title.font.size = Pt(16)
    
    run_sub_title = p_title.add_run(
        "Title: SMART INTERACTIVE POSITIVE AIRWAY PRESSURE SYSTEM (SIPAP) "
        "INTEGRATING REAL-TIME PRESSURE AND FLOW RATE MONITORING"
    )
    run_sub_title.bold = True
    run_sub_title.font.size = Pt(14)
    
    # Helper to add headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Dark Navy
        return p
        
    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x33, 0x66, 0x99) # Medium Blue
        return p
        
    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        return p

    # 5. Section 1: Authors and Advisor Information
    add_heading_1("I. GENERAL INFORMATION AND SCIENTIFIC ADVISOR")
    
    p = doc.add_paragraph()
    p.add_run("1. Scientific Advisor & Professional Director: ").bold = True
    p.add_run("Prof. Dr. Sc. Duong Quy Sy ")
    p.add_run("(President of the Vietnam Society of Sleep Medicine - VSSM).\n").italic = True
    
    p.add_run("2. Academic Sponsor and Professional Partner: ").bold = True
    p.add_run("Vietnam Society of Sleep Medicine (VSSM).\n").italic = True
    
    p.add_run("3. Research and Development Team: ").bold = True
    p.add_run(
        "Prof. Dr. Sc. Duong Quy Sy, MSc. Tran Huu Nam, MSc. Nguyen Tuan Anh, MSc. Tang Thi Thảo Trâm, "
        "BA. Tran Thi Cam Tu, Pharm. Nguyen Trong Bang, MSc. Nguyen Van Toi, PhD. Nguyen Duy Thai, "
        "Prof. Ramon Farre, Prof. Dr. Thomas Penzel, Prof. Dr. Francis Martin, Prof. Dinh Xuan Anh Tuan."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "The patented name of the device is SIPAP (Smart Interactive PAP - Smart Interactive Positive Airway Pressure Device), "
        "which is a brand term and structural solution exclusively created and proposed for intellectual property protection by Prof. Dr. Sc. Duong Quy Sy, "
        "to honor the professional partnership with the Vietnam Society of Sleep Medicine (VSSM)."
    ).italic = True

    add_image_with_caption(doc, "image_3.png", "Figure 1: Prof. Dr. Sc. Duong Quy Sy inspects the physical design prototype of the SIPAP smart interactive breathing support device.", 4.0)

    # 6. Section 2: Technical Field
    add_heading_1("II. TECHNICAL FIELD")
    p = doc.add_paragraph()
    p.add_run(
        "This utility solution refers to the field of respiratory support medical devices, specifically a "
        "Continuous Positive Airway Pressure (CPAP) system with smart interactive capability "
        "(Smart & Interactive PAP - abbreviated as SIPAP). This system serves the purpose of non-invasive respiratory support "
        "for the treatment and management at home or in hospitals of patients with Obstructive Sleep Apnea (OSA)."
    )

    # 7. Section 3: Technical Background (Prior Art)
    add_heading_1("III. TECHNICAL BACKGROUND OF THE UTILITY SOLUTION")
    p = doc.add_paragraph()
    p.add_run(
        "According to recent medical statistics, there are about 4 million adults in Vietnam suffering from "
        "Obstructive Sleep Apnea (OSA), accounting for approximately 8.5% of the population. Among them, there are "
        "about 2.3 million patients at moderate to severe levels - a high-risk group for dangerous cardiovascular "
        "complications (hypertension, stroke, myocardial infarction) and occupational accidents due to daytime sleepiness. "
        "The standard and most effective treatment method is using a continuous positive airway pressure (CPAP) device during sleep to maintain airway patency."
    )
    
    p = doc.add_paragraph()
    p.add_run("However, current commercial CPAP devices in Vietnam face the following major drawbacks:\n").bold = True
    add_bullet("Very high import costs (ranging from 15 to over 50 million VND), exceeding the financial capacity of most working-class people.")
    add_bullet("Closed system design, making it impossible to customize hardware or upgrade additional sensor modules tailored to clinical needs.")
    add_bullet("Lack of real-time direct connection with personal mobile devices and medical data synchronization with the remote health monitoring system (SleepCare), making it difficult for physicians to monitor patient compliance at home.")
    add_bullet("Some low-cost models do not integrate high-sensitivity respiratory flow sensors (Flow Sensor) or blood oxygen saturation sensors (SpO2) directly on the same system to synchronously calculate the patient's Apnea-Hypopnea Index (AHI).")

    # 8. Section 4: Summary of the Invention
    add_heading_1("IV. TECHNICAL NATURE OF THE UTILITY SOLUTION")
    p = doc.add_paragraph()
    p.add_run(
        "To resolve the aforementioned limitations, this utility solution proposes a smart CPAP system (SIPAP) "
        "with a highly modularized structure, integrating multi-sensor synchronization and Bluetooth Low Energy (BLE) wireless communication. "
        "The system is designed to optimize manufacturing costs, utilizing suitable components while still ensuring medical accuracy."
    )
    
    add_heading_2("1. Key Features and Operating Modes:")
    add_bullet("The system supports 3 basic software-customizable breathing modes: CPAP/AutoCPAP (fixed or automatically adjusted positive airway pressure), BiPAP (providing two distinct pressure levels: IPAP for inhalation and EPAP for exhalation to reduce work of breathing), and APAP (real-time airway obstruction analysis to adjust pressure optimally).")
    add_bullet("Stable therapeutic pressure range: 4 to 25 cmH2O, continuously regulated based on a closed-loop fan speed (RPM) control algorithm.")
    add_bullet("Integrated MPXV5010G pressure sensor to measure actual airway pressure at the mask with high sensitivity.")
    add_bullet("Integrated SFM3300-D flow sensor to measure instantaneous respiratory flow rate of the patient to determine respiratory rate, detect apnea or hypopnea events.")
    add_bullet("Pre-designed connection port for blood oxygen saturation (SpO2) sensor to facilitate integration in future upgrades.")
    add_bullet("Intuitive local user interface via OLED display screen for parameters and a rotary encoder combined with a push button.")
    add_bullet("BLE GATT Profile wireless connection with device name 'SIPAP' transmitting JSON data packets every 200ms to a mobile app for SleepCare cloud synchronization.")
    add_bullet("Integrated humidifier module containing distilled water in the airway path to maintain appropriate humidity and temperature for the air supplied to the patient, reducing irritation and dryness of the user's respiratory mucosa.")

    add_image_with_caption(doc, "image_en.png", "Figure 2: Operation flowchart of the SIPAP breathing support system with CPAP, BiPAP, and APAP modes.", 5.5)

    # 9. Section 5: Detailed Description of Implementation
    add_heading_1("V. DETAILED DESCRIPTION OF PREFERRED EMBODIMENTS")
    
    add_heading_2("1. System Hardware Block Diagram")
    p = doc.add_paragraph()
    p.add_run(
        "The system comprises the following key functional blocks connected and operating coordinately as follows:\n"
    )
    add_bullet("Vital sign data acquisition sensor block: MPXV5010G pressure sensor and SFM3300-D flow sensor are the two primary active sensors in version 1 to measure mask pressure and actual breathing flow rate. Other auxiliary sensors such as SHT30/SHT31 air temperature & humidity sensor, SpO2 sensor are pre-designed with standby ports for future upgrades.")
    add_bullet("Central processing block: The main microcontroller is Arduino Nano 33 BLE Sense (using a powerful 32-bit ARM Cortex-M4 chip, integrated BLE antenna) which performs sensor reading, runs the PID fan control algorithm, and broadcasts BLE. An auxiliary ESP32 microcontroller is redundantly integrated on-board to enable Wi-Fi/Internet connectivity for direct data transmission to the IoT Cloud.")
    add_bullet("Pressure generation block: Utilizing a WS4540-12-NZ03 brushless blower motor operating at 12VDC, along with a dedicated motor driver. The driver receives PWM pulses from pin D3 of Arduino to adjust rotational speed, and feeds back rotational speed pulses (FG) to external interrupt pin D2 of Arduino.")
    add_bullet("Airway block: Comprising an air chamber made of 3mm-thick transparent acrylic (Mica) sheets cut with high laser precision, medical-grade soft silicone tubing, a nasal mask, and an exhalation valve.")
    add_bullet("Humidification block (Humidifier): Integrating a humidification chamber filled with medical water located sequentially after the blower and before the nasal mask. Positive airway pressure air passes over the water surface to naturally absorb diffused moisture before entering the user's airway.")
    add_bullet("Power supply block: A high-power 12VDC adapter supplying the blower; a highly stable step-down Buck Converter circuit to 5V and 3.3V to power the microcontroller and sensors, preventing analog signal noise.")
    add_bullet("Interface block: OLED screen combined with a rotary encoder (with button switch) for local user configuration of target pressure and breathing mode.")

    add_image_with_caption(doc, "schematic_en.png", "Figure 3: Hardware schematic diagram and pin allocation between components of the SIPAP system.", 5.5)
    add_image_with_caption(doc, "image_4.png", "Figure 4: Close-up of the actual SIPAP device showing air chamber structure and board layout.", 4.0)
    add_image_with_caption(doc, "image_1.png", "Figure 5: Complete prototype of the SIPAP device connected to a breathing bag (artificial lung) for blower performance testing.", 4.0)

    add_heading_2("1.1. Detailed Technical Specifications of Core Components")
    
    p = doc.add_paragraph()
    p.add_run("a) Brushless Blower Motor (WS4540-12-NZ03):\n").bold = True
    p.add_run("This is the core component of the pressure generation block, producing the flow rate and positive pressure needed for the airway. Operational characteristics include:\n")
    add_bullet("Part Number: WS4540-12-NZ03.")
    add_bullet("Rated Voltage: 12 VDC.")
    add_bullet("Operation at Max Air Flow: Rotational speed reaches 45,000 RPM; Current consumption 1.6 A; Power consumption 119.2 W; Air flow rate reaches 7.2 m³/h (equivalent to 4.23 CFM or 120 LPM); Measured noise level is 62 dBA.")
    add_bullet("Operation at Max Air Pressure: Rotational speed reaches 49,000 RPM; Current consumption 0.9 A; Power consumption 10.8 W; Static pressure reaches 5.0 kPa (equivalent to 51 cmH2O); Blocked flow noise level is 49 dBA.")
    
    add_image_with_caption(doc, "blower.png", "Figure 6: Brushless blower motor WS4540-12-NZ03 used in the pressure generation block.", 3.2)
    
    p = doc.add_paragraph()
    p.add_run("b) Airway Pressure Sensor (MPXV5010G):\n").bold = True
    p.add_run("Specialized silicon piezoresistive sensor for measuring airway pressure at the mask. Technical specifications include:\n")
    add_bullet("Pressure Measurement Range (POP): 0 to 10 kPa (equivalent to 0 to 1019.78 mmH2O or 0 to 100 cmH2O).")
    add_bullet("Supply Voltage (VS): 4.75 VDC to 5.25 VDC (Rated at 5.0 VDC).")
    add_bullet("Current Consumption (Io): Max 10 mAdc (Typical 5.0 mAdc).")
    add_bullet("Zero-pressure Output Voltage (Voff): 0.2 VDC (Typical) at 0 kPa pressure.")
    add_bullet("Full Scale Output Voltage (VFSO): 4.7 VDC (Typical) at 10 kPa pressure.")
    add_bullet("Full Scale Span (VFSS): 4.5 VDC (Typical).")
    add_bullet("Accuracy: Error within ±5.0% VFSS in the operating temperature range of 0 to 85°C.")
    add_bullet("Sensitivity: 450 mV/kPa (or 4.413 mV/mmH2O).")
    add_bullet("Response Time: 1.0 ms.")
    add_bullet("Warm-Up Time: 20 ms.")
    add_bullet("Maximum Pressure Limit (Pmax): 40 kPa.")
    
    add_image_with_caption(doc, "sensorpressure.png", "Figure 7: Silicon pressure sensor MPXV5010G with SOP package layout.", 3.0)
    
    p = doc.add_paragraph()
    p.add_run("c) Respiratory Flow Sensor (Sensirion SFM3300-D / SFM3300-AW):\n").bold = True
    p.add_run("Specialized digital flow meter measuring medical respiratory gas flow. Technical specifications include:\n")
    add_bullet("Flow Measurement Range: -250 to +250 slm (Bi-directional).")
    add_bullet("Dead Space: Extremely small, under 10 ml.")
    add_bullet("Output Resolution: 14-bit.")
    add_bullet("Data Update Cycle (Update Time): Extremely fast, only 0.5 ms.")
    add_bullet("Communication Interface: I2C protocol, default sensor address: 64 (0x40 in hexadecimal).")
    add_bullet("Operating Voltage: 5 VDC ±5%.")
    add_bullet("Power Consumption: Under 50 mW.")
    add_bullet("Scale Factor Flow: 120 slm⁻¹.")
    add_bullet("Zero Offset Flow (Offset Flow): 32768.")
    add_bullet("Absolute Operating Pressure Range: 0.54 to 1.1 bar.")
    
    add_image_with_caption(doc, "sensorflow.png", "Figure 8: Sensirion SFM3300-D specialized medical gas flow sensor.", 3.5)

    # 10. Table: Bill of Materials (BOM)
    add_heading_2("2. Bill of Materials (BOM) for Prototype Fabrication")
    
    p = doc.add_paragraph()
    p.add_run(
        "Below is the list of materials and components for fabricating the SIPAP smart interactive breathing support system, "
        "compiled for the proposed prototype version:"
    )

    # Columns: No., Component Name, Qty, Unit Price (VND), Total Cost (VND), Note
    bom_data = [
        ("1", "Arduino Nano 33 BLE Sense", "1", "863.000", "863.000", "Main MCU, ARM Cortex-M4, integrated BLE"),
        ("2", "ESP32 Microcontroller (unused)", "1", "162.000", "162.000", "Standby MCU for Wi-Fi expansion (future upgrade)"),
        ("3", "12V Power Adapter", "1", "100.000", "100.000", "Supplies 12VDC power for blower and driver"),
        ("4", "DC Power Jack Adapter", "1", "50.000", "50.000", "5.5x2.1mm DC power jack adapter connector"),
        ("5", "Step-Down Buck Converter", "1", "50.000", "50.000", "Steps down 12V to 5V/3.3V for MCU and sensors"),
        ("6", "Cables, Wires, and Adhesives", "1", "50.000", "50.000", "Materials for electrical connections and housing assembly"),
        ("7", "Pressure Sensor (MPXV5010G)", "1", "673.000", "673.000", "Silicon pressure sensor for mask pressure (0-10 kPa)"),
        ("8", "Blower Motor & Driver (WS4540-12-NZ03)", "1", "1.142.000", "1.142.000", "High-pressure 12VDC blower motor and PWM driver board"),
        ("9", "Nasal Mask", "1", "1.000.000", "1.000.000", "Medical-grade silicone nasal mask for patient use"),
        ("10", "Acrylic Sheets + Laser Cutting Cost", "1", "500.000", "500.000", "For fabricating the device chassis and pressure chamber"),
        ("11", "SpO2 Sensor (unused)", "1", "100.000", "100.000", "Optical blood oxygen saturation sensor (future upgrade)"),
        ("12", "OLED Display", "1", "65.000", "65.000", "Displays pressure, flow rate, and respiratory rate"),
        ("13", "Flow Sensor (SFM3300-D)", "1", "2.300.000", "2.300.000", "Specialized medical gas flow sensor, I2C interface"),
        ("14", "Rotary Encoder with Push Button", "1", "40.000", "40.000", "Rotary encoder module for local user interaction"),
        ("15", "Humidifier Chamber", "1", "250.000", "250.000", "Medical water chamber for passive breathing gas humidification")
    ]
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Shading Accent 1'
    
    # Format Table headers
    hdr_cells = table.rows[0].cells
    headers = ["No.", "Component Name", "Qty", "Unit Price (VND)", "Total Cost (VND)", "Note"]
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        set_cell_shading(hdr_cells[i], "1B365D") # Navy header
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(11)
            
    # Fill Table data
    for stt, name, qty, price, total, note in bom_data:
        row_cells = table.add_row().cells
        data = [stt, name, qty, price, total, note]
        for i, val_text in enumerate(data):
            row_cells[i].text = val_text
            set_cell_margins(row_cells[i], top=100, bottom=100, left=120, right=120)
            p = row_cells[i].paragraphs[0]
            # Alignment
            if i in [0, 2]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i in [3, 4]:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            for run in p.runs:
                run.font.size = Pt(10)
                
    # Add Total row
    row_cells = table.add_row().cells
    row_cells[0].text = "TOTAL MATERIAL AND COMPONENT COST"
    set_cell_margins(row_cells[0], top=120, bottom=120, left=150, right=150)
    set_cell_shading(row_cells[0], "F2F2F2")
    p = row_cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.bold = True
        run.font.size = Pt(11)
        
    # Merge cells for total title
    row_cells[0].merge(row_cells[1]).merge(row_cells[2]).merge(row_cells[3])
    
    row_cells[4].text = "7.345.000"
    set_cell_margins(row_cells[4], top=120, bottom=120, left=150, right=150)
    set_cell_shading(row_cells[4], "F2F2F2")
    p = row_cells[4].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.bold = True
        run.font.size = Pt(11)
        
    row_cells[5].text = "Includes VAT and acrylic laser-cutting costs"
    set_cell_margins(row_cells[5], top=120, bottom=120, left=150, right=150)
    set_cell_shading(row_cells[5], "F2F2F2")
    p = row_cells[5].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.italic = True
        run.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 11. Section 5.3: Firmware operation and Algorithms
    add_heading_2("3. Control Algorithms and Software Operation")
    p = doc.add_paragraph()
    p.add_run(
        "Real-time processing of the device is optimized and programmed on the microcontroller platform. "
        "Based on the actual firmware source code ('dull.ino'), the data processing principles are executed as follows:"
    )
    
    # Bullet points of algorithm
    add_bullet("Verification of the pressure calculation algorithm from the MPXV5010G sensor:\n"
               "The pressure reading function in the source code 'dull.ino' uses the formula to convert the ADC value into actual pressure (kPa):\n"
               "  `pressure = ((sensorVoltage / 5.0) - 0.04) / 0.09;`\n"
               "This equation is mathematically and physically correct, verified by the manufacturer's datasheet:\n"
               "  - According to the MPXV5010G datasheet, the theoretical output transfer function (with supply Vs = 5.0 Vdc, temperature TA = 25°C) is: Vout = Vs * (0.09 * P + 0.04), where P is the input pressure in kPa.\n"
               "  - Solving for P from output voltage Vout:\n"
               "    + Divide both sides by Vs: Vout / Vs = 0.09 * P + 0.04\n"
               "    + Subtract 0.04 from both sides: (Vout / Vs) - 0.04 = 0.09 * P\n"
               "    + Divide both sides by 0.09: P = ((Vout / Vs) - 0.04) / 0.09 (kPa)\n"
               "  - Code alignment: The variable `sensorVoltage` represents Vout. The float literal `5.0` represents supply voltage Vs = 5.0V. The variable `sensorVoltage` is reconstructed from the voltage read at ADC pin A0 of the microcontroller using a scale factor of 1.5 (`sensorVoltage = voltage * 1.5`) to compensate for the external 10k/20k voltage divider bridge (which reduces output voltage by 1.5 times, meaning V_adc = Vout * (2/3) => Vout = V_adc * 1.5). This voltage divider is used to limit the 5V sensor output to prevent damaging the 3.3V ADC input pin of the Arduino Nano 33 BLE.\n"
               "Thus, the formula in the source code is entirely correct, faithfully reflecting the physical pressure measured at the patient's mask.\n"
               "Similarly, for the Sensirion SFM3300-D digital flow sensor, the formula converting the 14-bit raw I2C value (Raw_Flow) to actual flow rate (slm) is:\n"
               "  `Flow (slm) = (Raw_Flow - 32768) / 120.0`\n"
               "This formula is entirely correct according to the Sensirion datasheet, with a zero-offset Flow of 32768 and a Scale Factor Flow of 120 slm⁻¹.")
               
    add_bullet("Measuring blower motor rotational speed (RPM):\n"
               "To control the WS4540 blower motor accurately and stably, its speed feedback pulse (FG) is connected to Arduino pin D2 "
               "using an external interrupt triggered on the rising edge (RISING). The interrupt service routine countPulse() increments the count: \n"
               "  - pulseCount = pulseCount + 1\n"
               "Every 1000ms (1 second), the microcontroller briefly disables interrupts, reads pulseCount, resets pulseCount to 0, "
               "and re-enables interrupts. Since the WS4540 blower generates 2 FG pulses per rotation, the motor speed (RPM - Revolutions Per Minute) is calculated using the formula:\n"
               "  - RPM = (pulseCount / 2) * 60 (RPM).")
               
    add_bullet("Blower Speed Control Mechanism (Blower Control):\n"
               "Blower speed is regulated via a Pulse Width Modulation (PWM) signal written to pin D3 (analogWrite). Value ranges from 0 (fan stopped) to 255 (maximum speed). "
               "Specifically, the 'START' command sets a safe initial starting PWM value of 55, while the 'STOP' command sets the PWM to 0. "
               "The system also supports custom PWM values sent from mobile devices for target pressure fine-tuning.")
               
    add_bullet("BLE (Bluetooth Low Energy) Wireless Connection Protocol:\n"
               "The system operates as a BLE peripheral (BLE Peripheral). The detailed GATT Profile configuration includes:\n"
               "  - Broadcast Device Name: 'SIPAP' (or 'CPAP_VSSM')\n"
               "  - CPAP Service UUID: cb24858f-399f-4498-85e8-fea9d383d54f\n"
               "  - Sensor Characteristic: UUID 5e9e214b-124c-434d-84e5-018dccd35df1 (Supports Read and Notify). Periodically transmits a standard JSON status package every 200ms: "
               '{"pressure": <value>, "rpm": <value>, "pwm": <value>} to update the mobile application in real-time.\n'
               "  - Action Characteristic: UUID 56debc28-acab-4184-8f86-1a9c887b220a (Supports Write). Receives string command control sequences from the mobile application: "
               "'LED_ON' (turns indicator LED on), 'LED_OFF' (turns indicator LED off), 'START' (starts blower), 'STOP' (stops blower), 'PWM:<value>' (sets blower speed directly).\n"
               "  - Safety Feature: Upon BLE disconnection (disconnect), the device automatically stops the blower (PWM = 0) to protect the patient.")

    add_image_with_caption(doc, "image_2.png", "Figure 9: Experimental testing measuring parameters and controlling blower speed on a patient simulator model connected to a PC.", 4.0)

    # 12. Section 6: Technical and Economic Benefits
    add_heading_1("VI. ACHIEVED BENEFITS AND FUTURE ROADMAP")
    p = doc.add_paragraph()
    p.add_run("1. Technical and Economic Benefits:\n").bold = True
    add_bullet("Extremely low manufacturing cost: The total component cost is only about 7,345,000 VND, which is 3 to 8 times cheaper than commercial imported CPAP machines in the market.")
    add_bullet("Integrated sensors for real-time airway pressure and flow rate measurement, providing crucial respiratory data for the patient during sleep.")
    add_bullet("High IoT potential: Facilitated by BLE connectivity and the integrated ESP32 module, breathing data can be easily synchronized with the SleepCare medical cloud, enabling remote diagnosis and monitoring by physicians.")
    add_bullet("Compact design and laser-cut transparent acrylic housing make visual mechanical checks easy, allowing quick modular installation and repairs.")

    p = doc.add_paragraph()
    p.add_run("2. Strategic Development Roadmap (Under VSSM Clinical Research):\n").bold = True
    add_bullet("Phase 1: Clinical Trial & Safety Certifications - Conduct clinical trials of the prototype to evaluate medical safety and calibrate the optimal pressure range under the supervision of Prof. Dr. Sc. Duong Quy Sy.")
    add_bullet("Phase 2: Cloud Integration & Remote Monitoring - Finalize firmware for the secondary ESP32 microcontroller to support direct Wi-Fi connectivity, uploading data directly to the cloud server without a smartphone intermediary.")
    add_bullet("Phase 3: AI-based Adaptation - Develop and embed artificial intelligence (AI/Machine Learning) algorithms on the mobile app or server to automatically identify abnormal breathing patterns, predict apnea events early, and dynamically adjust blower pressure (Auto-CPAP/APAP) in a smarter way.")
    add_bullet("Phase 4: Sleep Ecosystem Integration - Integrate the device into a comprehensive smart sleep care ecosystem along with environmental sensors (bedroom temperature, light, noise).")

    add_image_with_caption(doc, "image5_en.png", "Figure 10: Development roadmap of the SIPAP system from MVP to a smart sleep care ecosystem (simulation model diagram).", 6.0)

    # 13. Section 7: Patent Claims (Yêu cầu bảo hộ)
    add_heading_1("VII. PATENT CLAIMS")
    p = doc.add_paragraph()
    p.add_run("1. A smart interactive continuous positive airway pressure (CPAP) system (SIPAP) for supporting treatment of obstructive sleep apnea syndrome, the system comprising:")
    add_bullet("A positive pressure generation block comprising a 12VDC brushless blower motor and a PWM-based speed controller;")
    add_bullet("An integrated measurement sensor block comprising an analog pressure sensor connected via a voltage divider bridge to monitor nasal mask pressure, and a respiratory flow sensor communicating via standard I2C protocol;")
    add_bullet("A central microcontroller block incorporating a Bluetooth Low Energy (BLE) communication module, electrically connected to the pressure generation block and the sensor block;")
    add_bullet("characterized in that: the central microcontroller is programmed to continuously read nasal mask pressure signals from the pressure sensor, calculate flow rate from the flow sensor, and simultaneously measure blower speed (RPM) by counting speed feedback pulses connected to an external interrupt pin of the microcontroller, and periodically transmit a JSON-formatted data string containing pressure, RPM, and blower PWM speed values via BLE GATT Profile under the broadcast name 'SIPAP' for real-time monitoring.")
    
    p = doc.add_paragraph()
    p.add_run("2. The breathing support system according to claim 1, characterized in that: ")
    p.add_run("the hardware is designed to incorporate a standby connection port and the source code is pre-configured with redundant logic to integrate an optical blood oxygen saturation (SpO2) sensor in a subsequent upgrade.")
    
    p = doc.add_paragraph()
    p.add_run("3. The breathing support system according to claim 1 or 2, characterized in that: ")
    p.add_run("the central microcontroller is configured with an automatic safety protection feature to immediately turn off the blower (setting PWM to 0) upon detecting a connection loss of the wireless BLE communication with the mobile application.")
    
    p = doc.add_paragraph()
    p.add_run("4. The breathing support system according to claim 1, characterized in that: ")
    p.add_run("the housing of the breathing device and the blower pressure chamber are fabricated monolithically from transparent acrylic (Mica) sheets of 3-5mm thickness cut by laser to ensure airtightness and direct visual inspection of internal hardware components.")
    
    p = doc.add_paragraph()
    p.add_run("5. The breathing support system according to claim 1, characterized in that: ")
    p.add_run("the system integrates a humidification chamber (Humidifier) containing medical distilled water positioned sequentially downstream from the positive pressure generation block and upstream from the nasal mask, to naturally humidify the passing air flow, preventing irritation and dryness of the user's respiratory mucosa.")
    
    p = doc.add_paragraph()
    p.add_run("6. The breathing support system according to claim 1, characterized in that: ")
    p.add_run("the exclusive patent protects the name 'SIPAP' (Smart Interactive PAP - Smart Interactive Positive Airway Pressure Device) proposed by Prof. Dr. Sc. Duong Quy Sy, associated with the medical technical system solution ecosystem described above.")

    # Save to file
    out_dir = r"c:\Users\ADMIN\OneDrive\Máy tính\Master_2024\Cpap\document"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "utility_solution_SIPAP_v1_en.docx")
    doc.save(out_path)
    print(f"Document saved successfully at {out_path}!")

if __name__ == "__main__":
    create_document()
