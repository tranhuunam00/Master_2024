import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def add_image_with_caption(doc, img_name, caption_text, width_inches):
    img_path = os.path.join(r"c:\Users\ADMIN\OneDrive\Máy tính\Master_2024\Cpap\document", img_name)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(img_path, width=Inches(width_inches))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(8)
        run_cap = p_cap.add_run(caption_text)
        run_cap.italic = True
        run_cap.font.size = Pt(10)
    else:
        print(f"Warning: Image {img_name} not found at {img_path}")

def create_document():
    doc = Document()
    
    # 1. Page Margins Setup (Standard Vietnamese administrative layout: Left 3cm, Right 2cm, Top 2cm, Bottom 2cm)
    # 1 inch = 2.54 cm -> 3cm = 1.18 in, 2cm = 0.79 in
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.79)
        
    # 2. Typography and Styles Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    # Custom Paragraph format for Normal style
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # 3. Header Administrative Block
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header_format = p_header.paragraph_format
    p_header_format.space_after = Pt(2)
    
    run_nation1 = p_header.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    run_nation1.bold = True
    run_nation1.font.size = Pt(12)
    
    run_nation2 = p_header.add_run("Độc lập - Tự do - Hạnh phúc\n")
    run_nation2.bold = True
    run_nation2.font.size = Pt(13)
    
    run_line = p_header.add_run("---------------o0o---------------")
    run_line.font.size = Pt(10)
    
    # 4. Title of the Patent
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(18)
    p_title.paragraph_format.space_after = Pt(18)
    
    run_title = p_title.add_run("BẢN MÔ TẢ GIẢI PHÁP HỮU ÍCH\n")
    run_title.bold = True
    run_title.font.size = Pt(16)
    
    run_sub_title = p_title.add_run(
        "Tên giải pháp: HỆ THỐNG HỖ TRỢ THỞ ÁP LỰC DƯƠNG TƯƠNG TÁC THÔNG MINH (SIPAP)\n"
        "TÍCH HỢP GIÁM SÁT ÁP SUẤT VÀ LƯỢNG KHÍ THỞ THỜI GIAN THỰC"
    )
    run_sub_title.bold = True
    run_sub_title.font.size = Pt(14)
    
    # Helper to add headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Dark Navy
        return p
        
    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x33, 0x66, 0x99) # Medium Blue
        return p
        
    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        return p

    # 5. Section 1: Authors and Advisor Information
    add_heading_1("I. THÔNG TIN CHUNG VÀ CỐ VẤN KHOA HỌC")
    
    p = doc.add_paragraph()
    p.add_run("1. Cố vấn khoa học & Chỉ đạo chuyên môn: ").bold = True
    p.add_run("GS. TS. KH. Dương Quý Sỹ ")
    p.add_run("(Chủ tịch Hội Y học Giấc ngủ Việt Nam - VSSM).\n").italic = True
    
    p.add_run("2. Đơn vị bảo trợ học thuật và đồng hành chuyên môn: ").bold = True
    p.add_run("Hội Y học Giấc ngủ Việt Nam (VSSM).\n").italic = True
    
    p.add_run("3. Nhóm tác giả nghiên cứu chế tạo: ").bold = True
    p.add_run(
        "GS.TS.KH. Dương Quý Sỹ, ThS. Trần Hữu Nam, ThS. Nguyễn Tuấn Anh, ThS. Tăng Thị Thảo Trâm, "
        "CN. Trần Thị Cẩm Tú, DS. Nguyễn Trọng Bằng, ThS. Nguyễn Văn Tới, TS. Nguyễn Duy Thái, "
        "GS. Ramon Farré, GS.TS. Thomas Penzel, GS.TS. Francis Martin, GS. Đinh Xuân Anh Tuấn."
    )
    
    p = doc.add_paragraph()
    p.add_run(
        "Tên gọi sáng chế của thiết bị là SIPAP (Smart Interactive PAP - Thiết bị hỗ trợ thở áp lực dương tương tác thông minh), "
        "là thuật ngữ thương hiệu và giải pháp cấu trúc do GS.TS.KH. Dương Quý Sỹ độc quyền sáng tạo và đề xuất bảo hộ quyền sở hữu trí tuệ, "
        "nhằm vinh danh sự đồng hành chuyên môn từ Hội Y học Giấc ngủ Việt Nam."
    ).italic = True

    add_image_with_caption(doc, "image (3).png", "Hình 1: GS. TS. KH. Dương Quý Sỹ kiểm tra mẫu thiết kế thực tế của thiết bị hỗ trợ thở tương tác thông minh SIPAP.", 4.0)

    # 6. Section 2: Technical Field
    add_heading_1("II. LĨNH VỰC KỸ THUẬT ĐƯỢC ĐỀ CẬP")
    p = doc.add_paragraph()
    p.add_run(
        "Giải pháp hữu ích này đề cập đến lĩnh vực thiết bị y tế hỗ trợ chức năng hô hấp, cụ thể là "
        "hệ thống hỗ trợ thở áp lực dương liên tục (Continuous Positive Airway Pressure - CPAP) có khả năng tương tác thông minh "
        "(Smart & Interactive PAP - viết tắt là SIPAP). Hệ thống này phục vụ mục đích hỗ trợ thở không xâm lấn nhằm điều trị "
        "và quản lý tại nhà hoặc bệnh viện cho các bệnh nhân mắc hội chứng ngưng thở khi ngủ do tắc nghẽn (Obstructive Sleep Apnea - OSA)."
    )

    # 7. Section 3: Technical Background (Prior Art)
    add_heading_1("III. TÌNH TRẠNG KỸ THUẬT CỦA GIẢI PHÁP HỮU ÍCH")
    p = doc.add_paragraph()
    p.add_run(
        "Theo các thống kê y tế gần đây, tại Việt Nam có khoảng 4 triệu người trưởng thành mắc hội chứng ngưng thở khi ngủ do tắc nghẽn (OSA), "
        "chiếm tỷ lệ khoảng 8.5% dân số. Trong đó, có khoảng 2.3 triệu bệnh nhân ở mức độ trung bình đến nặng - là nhóm đối tượng có nguy cơ cao "
        "gặp các biến chứng nguy hiểm về tim mạch (tăng huyết áp, đột quỵ, nhồi máu cơ tim) và tai nạn lao động do buồn ngủ ban ngày. "
        "Phương pháp điều trị chuẩn và hiệu quả nhất là sử dụng máy thở áp lực dương liên tục (CPAP) trong lúc ngủ để duy trì đường thở luôn thông thoáng."
    )
    
    p = doc.add_paragraph()
    p.add_run("Tuy nhiên, các máy CPAP thương mại hiện hành tại Việt Nam đối mặt với các nhược điểm lớn sau:\n").bold = True
    add_bullet("Chi phí nhập khẩu rất cao (dao động từ 15 đến hơn 50 triệu đồng), vượt quá khả năng chi trả của phần lớn người dân lao động.")
    add_bullet("Hệ thống đóng kín, không thể cá nhân hóa phần cứng hoặc nâng cấp các module cảm biến bổ sung tùy biến theo nhu cầu lâm sàng.")
    add_bullet("Thiếu khả năng kết nối trực tiếp thời gian thực với các thiết bị di động cá nhân và đồng bộ dữ liệu y tế với hệ thống theo dõi sức khỏe từ xa (SleepCare), gây khó khăn cho các bác sĩ trong việc giám sát sự tuân thủ điều trị của bệnh nhân tại nhà.")
    add_bullet("Một số dòng máy giá rẻ không tích hợp sẵn cảm biến lưu lượng khí thở độ nhạy cao (Flow Sensor) hoặc cảm biến bão hòa oxy máu (SpO2) trực tiếp trên cùng một hệ thống để tính toán đồng bộ chỉ số giảm thở AHI (Apnea-Hypopnea Index) của bệnh nhân.")

    # 8. Section 4: Summary of the Invention
    add_heading_1("IV. BẢN CHẤT KỸ THUẬT CỦA GIẢI PHÁP HỮU ÍCH")
    p = doc.add_paragraph()
    p.add_run(
        "Để giải quyết các hạn chế nêu trên, giải pháp hữu ích đề xuất một hệ thống CPAP thông minh (SIPAP) "
        "có cấu trúc module hóa cao, tích hợp đồng bộ đa cảm biến và giao tiếp không dây năng lượng thấp (Bluetooth Low Energy - BLE). "
        "Hệ thống được thiết kế hướng tới tối ưu hóa chi phí sản xuất, sử dụng các linh kiện phù hợp nhưng vẫn đảm bảo độ chính xác y tế."
    )
    
    add_heading_2("1. Các tính năng và chế độ hoạt động chính:")
    add_bullet("Hệ thống hỗ trợ 3 chế độ thở cơ bản tùy biến qua phần mềm: CPAP/AutoCPAP (áp lực dương cố định hoặc tự động điều chỉnh áp lực), BiPAP (cung cấp hai mức áp lực riêng biệt IPAP khi hít vào và EPAP khi thở ra để giảm công hô hấp) và APAP (phân tích tắc nghẽn đường thở thời gian thực để thay đổi áp suất tối ưu).")
    add_bullet("Dải áp lực điều trị ổn định: từ 4 đến 25 cmH2O, được điều tiết liên tục dựa trên thuật toán điều khiển vòng kín tốc độ quạt (RPM).")
    add_bullet("Tích hợp cảm biến áp suất MPXV5010G đo áp suất đường thở thực tế tại vị trí mặt nạ với độ nhạy cao.")
    add_bullet("Tích hợp cảm biến lưu lượng SFM3300-D đo lưu lượng dòng khí thở tức thời của bệnh nhân để xác định nhịp thở, phát hiện cơn ngừng thở hoặc giảm thở.")
    add_bullet("Thiết kế dự phòng sẵn cổng kết nối cho cảm biến đo nồng độ bão hòa oxy máu (SpO2) nhằm phục vụ việc nâng cấp tích hợp ở phiên bản tiếp theo.")
    add_bullet("Giao diện người dùng tại chỗ trực quan thông qua màn hình OLED hiển thị thông số và nút xoay mã hóa vòng quay (Rotary Encoder) kết hợp nút nhấn.")
    add_bullet("Kết nối không dây BLE GATT Profile với tên thiết bị 'SIPAP' truyền gói dữ liệu JSON mỗi 200ms lên ứng dụng di động để đồng bộ với đám mây SleepCare.")

    add_image_with_caption(doc, "image.png", "Hình 2: Sơ đồ luồng hoạt động của hệ thống hỗ trợ thở SIPAP với 3 chế độ CPAP, BiPAP và APAP.", 5.5)

    # 9. Section 5: Detailed Description of Implementation
    add_heading_1("V. MÔ TẢ CHI TIẾT PHƯƠNG ÁN THỰC HIỆN")
    
    add_heading_2("1. Sơ đồ cấu trúc phần cứng hệ thống")
    p = doc.add_paragraph()
    p.add_run(
        "Hệ thống bao gồm các khối chức năng chính được kết nối và phối hợp vận hành như sau:\n"
    )
    add_bullet("Khối cảm biến thu thập dữ liệu sinh hiệu: Cảm biến áp suất MPXV5010G và cảm biến lưu lượng SFM3300-D là hai cảm biến hoạt động chính trong phiên bản 1 để đo áp suất mặt nạ và lưu lượng khí thở thực tế. Các cảm biến phụ trợ khác như cảm biến nhiệt độ & độ ẩm không khí SHT30/SHT31, cảm biến SpO2 được thiết kế sẵn cổng kết nối chờ để phục vụ nâng cấp ở các phiên bản tiếp theo.")
    add_bullet("Khối xử lý trung tâm: Vi điều khiển chính Arduino Nano 33 BLE Sense (sử dụng chip ARM Cortex-M4 32-bit mạnh mẽ, tích hợp sẵn anten BLE) thực hiện việc đọc dữ liệu cảm biến, chạy thuật toán điều khiển PID quạt và phát sóng BLE. Vi điều khiển phụ ESP32 được dự phòng tích hợp trên bo mạch để thực hiện kết nối Wi-Fi/Internet gửi dữ liệu trực tiếp lên Cloud IoT.")
    add_bullet("Khối tạo áp lực: Sử dụng động cơ quạt thổi không chổi than (Brushless Blower) WS4540-12-NZ03 hoạt động ở điện áp 12VDC, kèm Driver điều khiển chuyên dụng. Driver nhận xung PWM từ chân D3 của Arduino để thay đổi tốc độ quay, đồng thời gửi lại xung phản hồi tốc độ quay FG về chân ngắt D2 của Arduino.")
    add_bullet("Khối đường dẫn khí: Bao gồm buồng khí chế tạo từ vật liệu Mica trong dày 3mm cắt laser chính xác, ống dẫn khí silicon y tế mềm, mặt nạ thở dạng mũi (nasal mask) và van thở ra (exhalation valve).")
    add_bullet("Khối nguồn: Adapter 12VDC công suất lớn cấp cho động cơ blower; mạch hạ áp Buck Converter bước xuống 5V và 3.3V có độ ổn định cao để nuôi vi điều khiển và các cảm biến tránh nhiễu tín hiệu analog.")
    add_bullet("Khối giao diện: Màn hình OLED kết hợp bộ mã hóa quay (Rotary Encoder) 8 nút xoay và nút ấn để người dùng cài đặt thông số áp suất đích và chế độ thở tại chỗ.")

    add_image_with_caption(doc, "schematic.png", "Hình 3: Sơ đồ kết nối phần cứng và phân bổ chân tín hiệu giữa các thành phần của hệ thống SIPAP.", 5.5)
    add_image_with_caption(doc, "image (4).png", "Hình 4: Thiết bị thực tế SIPAP cận cảnh cấu trúc buồng khí và cách bố trí các module bo mạch.", 4.0)
    add_image_with_caption(doc, "image (1).png", "Hình 5: Mẫu thử nghiệm hoàn chỉnh của thiết bị SIPAP kết nối với bóng thở (phổi giả) phục vụ kiểm định hiệu năng quạt thổi.", 4.0)

    add_heading_2("1.1. Chi tiết thông số kỹ thuật của các cấu phần cốt lõi")
    
    p = doc.add_paragraph()
    p.add_run("a) Động cơ quạt thổi không chổi than (Brushless Blower WS4540-12-NZ03):\n").bold = True
    p.add_run("Đây là bộ phận cốt lõi của khối tạo áp lực, tạo ra lưu lượng và áp suất dương cần thiết cho đường dẫn khí. Các thông số đặc tính vận hành bao gồm:\n")
    add_bullet("Mã thiết bị (Part No): WS4540-12-NZ03.")
    add_bullet("Điện áp định mức (Voltage): 12 VDC.")
    add_bullet("Vận hành tại điểm Lưu lượng khí cực đại (At Max Air flow): Tốc độ quay đạt 45.000 vòng/phút (RPM); Dòng điện tiêu thụ 1,6 A; Công suất tiêu thụ 119,2 W; Lưu lượng khí đạt 7,2 m³/h (tương đương 4,23 CFM hoặc 120 LPM); Độ ồn đo được là 62 dBA.")
    add_bullet("Vận hành tại điểm Á suất khí cực đại (At Max Air Pressure): Tốc độ quay đạt 49.000 vòng/phút (RPM); Dòng điện tiêu thụ 0,9 A; Công suất tiêu thụ 10,8 W; Áp suất tĩnh đạt 5,0 kPa (tương đương 51 cmH2O); Độ ồn khi bị chặn hoàn toàn dòng khí (Block Noise Level) là 49 dBA.")
    
    add_image_with_caption(doc, "blower.png", "Hình 6: Động cơ quạt thổi không chổi than WS4540-12-NZ03 sử dụng trong khối tạo áp lực.", 3.2)
    
    p = doc.add_paragraph()
    p.add_run("b) Cảm biến áp suất đường thở (MPXV5010G):\n").bold = True
    p.add_run("Cảm biến silicon piezoresistive chuyên dụng để đo áp suất đường thở tại mặt nạ thở. Thông số đặc tính kỹ thuật bao gồm:\n")
    add_bullet("Dải đo áp suất (POP): 0 đến 10 kPa (tương đương 0 đến 1019,78 mmH2O hoặc 0 đến 100 cmH2O).")
    add_bullet("Điện áp cấp (VS): 4,75 VDC đến 5,25 VDC (Mức định mức 5,0 VDC).")
    add_bullet("Dòng điện tiêu thụ (Io): tối đa 10 mAdc (điển hình 5,0 mAdc).")
    add_bullet("Điện áp điểm không (Voff): 0,2 VDC (điển hình) tại áp suất 0 kPa.")
    add_bullet("Điện áp đầu ra cực đại (VFSO): 4,7 VDC (điển hình) tại áp suất 10 kPa.")
    add_bullet("Hiệu số điện áp toàn dải (VFSS): 4,5 VDC (điển hình).")
    add_bullet("Độ chính xác (Accuracy): sai số trong khoảng ±5,0% VFSS trong dải nhiệt độ vận hành từ 0 đến 85°C.")
    add_bullet("Độ nhạy (Sensitivity): 450 mV/kPa (hoặc 4,413 mV/mmH2O).")
    add_bullet("Thời gian đáp ứng (Response Time): 1,0 ms.")
    add_bullet("Thời gian khởi động ổn định (Warm-Up Time): 20 ms.")
    add_bullet("Áp suất chịu đựng tối đa trước khi hỏng (Pmax): 40 kPa.")
    
    add_image_with_caption(doc, "sensorpressure.png", "Hình 7: Cảm biến áp suất silicon MPXV5010G với kết cấu chân dán SOP.", 3.0)
    
    p = doc.add_paragraph()
    p.add_run("c) Cảm biến lưu lượng khí thở (Sensirion SFM3300-D / SFM3300-AW):\n").bold = True
    p.add_run("Lưu lượng kế kỹ thuật số chuyên dụng đo dòng khí hô hấp y tế. Thông số đặc tính kỹ thuật bao gồm:\n")
    add_bullet("Dải đo lưu lượng: -250 đến +250 slm (Đo hai chiều).")
    add_bullet("Khoảng chết (Dead space): cực nhỏ, dưới 10 ml.")
    add_bullet("Độ phân giải đầu ra: 14 bit.")
    add_bullet("Chu kỳ cập nhật dữ liệu (Update Time): cực nhanh, chỉ 0,5 ms.")
    add_bullet("Giao tiếp truyền thông: Giao thức I2C, Địa chỉ mặc định (Default Sensor Address): 64 (0x40 ở hệ cơ số 16).")
    add_bullet("Điện áp hoạt động: 5 VDC ±5%.")
    add_bullet("Công suất tiêu thụ: dưới 50 mW.")
    add_bullet("Hệ số tỉ lệ dòng (Scale Factor Flow): 120 slm⁻¹.")
    add_bullet("Giá trị điểm không (Offset Flow): 32768.")
    add_bullet("Dải áp suất tuyệt đối hoạt động: 0,54 đến 1,1 bar.")
    
    add_image_with_caption(doc, "sensorflow.png", "Hình 8: Cảm biến lưu lượng khí y tế chuyên dụng Sensirion SFM3300-D.", 3.5)

    add_heading_2("2. Danh mục vật tư chế tạo mẫu thử nghiệm (BOM)")
    
    p = doc.add_paragraph()
    p.add_run(
        "Dưới đây là bảng liệt kê danh mục vật tư chế tạo hệ thống máy thở tương tác thông minh SIPAP "
        "được lập cho phiên bản thử nghiệm - dự kiến:"
    )

    # 10. Table: Bill of Materials (BOM)
    # Columns: STT, Tên Linh Kiện, Số lượng, Đơn giá (VND), Thành tiền (VND), Ghi chú
    bom_data = [
        ("1", "Arduino 33 BLE SENSE", "1", "863.000", "863.000", "Vi điều khiển chính, ARM Cortex-M4, tích hợp BLE"),
        ("2", "ESP32 Microcontroller (chưa dùng)", "1", "162.000", "162.000", "Vi điều khiển phụ phục vụ mở rộng WiFi (dự phòng nâng cấp)"),
        ("3", "Nguồn 12 V", "1", "100.000", "100.000", "Cấp nguồn 12VDC cho quạt thổi và Driver"),
        ("4", "Rắc chuyển đổi", "1", "50.000", "50.000", "Đầu giắc cắm DC Power Jack 5.5x2.1mm"),
        ("5", "Mạch hạ áp (Buck)", "1", "50.000", "50.000", "Hạ áp từ 12V sang 5V/3.3V cấp cho MCU & cảm biến"),
        ("6", "Cáp nối, dây điện, keo dán", "1", "50.000", "50.000", "Vật tư kết nối mạch và cố định linh kiện cơ khí"),
        ("7", "Cảm biến áp suất (MPXV5010G)", "1", "673.000", "673.000", "Cảm biến áp suất silicon đo áp suất mask (0-10 kPa)"),
        ("8", "Motor Blower & Driver (WS4540-12-NZ03)", "1", "1.142.000", "1.142.000", "Động cơ thổi áp suất cao 12VDC và Driver PWM"),
        ("9", "Mặt nạ mũi (Nasal Mask)", "1", "1.000.000", "1.000.000", "Mặt nạ thở silicon y tế đeo mũi của bệnh nhân"),
        ("10", "Mica trong + công cắt laser", "1", "500.000", "500.000", "Chế tạo hộp vỏ máy và buồng nén khí"),
        ("11", "Cảm biến đo nồng độ SpO2 (chưa dùng)", "1", "100.000", "100.000", "Cảm biến quang học đo nồng độ oxy huyết (dự phòng nâng cấp cho phiên bản sau)"),
        ("12", "Màn hình LED hiển thị (OLED)", "1", "65.000", "65.000", "Màn hình hiển thị thông số áp suất, lưu lượng, nhịp thở"),
        ("13", "Cảm biến lưu lượng (SFM3300-D)", "1", "2.300.000", "2.300.000", "Cảm biến lưu lượng khí y tế chuyên dụng, giao tiếp I2C"),
        ("14", "Nút xoay và nút ấn (Rotary Encoder)", "8", "5.000", "40.000", "Bộ mã hóa vòng quay tương tác người dùng tại chỗ")
    ]
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Shading Accent 1'
    
    # Format Table headers
    hdr_cells = table.rows[0].cells
    headers = ["STT", "Tên Linh Kiện", "SL", "Đơn giá (VND)", "Thành tiền (VND)", "Ghi chú"]
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        set_cell_shading(hdr_cells[i], "1B365D") # Navy header
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(11)
            
    # Fill Table data
    for stt, name, qty, price, total, note in bom_data:
        row_cells = table.add_row().cells
        data = [stt, name, qty, price, total, note]
        for i, val_text in enumerate(data):
            row_cells[i].text = val_text
            set_cell_margins(row_cells[i], top=100, bottom=100, left=120, right=120)
            p = row_cells[i].paragraphs[0]
            # Alignment
            if i in [0, 2]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif i in [3, 4]:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            for run in p.runs:
                run.font.size = Pt(10)
                
    # Add Total row
    row_cells = table.add_row().cells
    row_cells[0].text = "TỔNG CỘNG CHI PHÍ VẬT TƯ CHẾ TẠO"
    set_cell_margins(row_cells[0], top=120, bottom=120, left=150, right=150)
    set_cell_shading(row_cells[0], "F2F2F2")
    p = row_cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.bold = True
        run.font.size = Pt(11)
        
    # Merge cells for total title
    row_cells[0].merge(row_cells[1]).merge(row_cells[2]).merge(row_cells[3])
    
    row_cells[4].text = "7.095.000"
    set_cell_margins(row_cells[4], top=120, bottom=120, left=150, right=150)
    set_cell_shading(row_cells[4], "F2F2F2")
    p = row_cells[4].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.bold = True
        run.font.size = Pt(11)
        
    row_cells[5].text = "Đã bao gồm VAT và công cắt vỏ mica"
    set_cell_margins(row_cells[5], top=120, bottom=120, left=150, right=150)
    set_cell_shading(row_cells[5], "F2F2F2")
    p = row_cells[5].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.italic = True
        run.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # 11. Section 5.3: Firmware operation and Algorithms
    add_heading_2("3. Giải thuật điều khiển và hoạt động của phần mềm")
    p = doc.add_paragraph()
    p.add_run(
        "Hoạt động xử lý thời gian thực của thiết bị được lập trình tối ưu hóa trên nền tảng vi điều khiển. "
        "Dựa trên mã nguồn firmware thực tế ('dull.ino'), nguyên lý xử lý số liệu được thực thi như sau:"
    )
    
    # Bullet points of algorithm
    add_bullet("Kiểm chứng thuật toán tính toán áp suất từ cảm biến MPXV5010G:\n"
               "Hàm đọc dữ liệu cảm biến áp suất trong mã nguồn 'dull.ino' sử dụng công thức chuyển đổi từ giá trị ADC sang giá trị áp suất thực tế (kPa):\n"
               "  `pressure = ((sensorVoltage / 5.0) - 0.04) / 0.09;`\n"
               "Phương trình này hoàn toàn chính xác về mặt toán học và vật lý học, được chứng minh dựa trên thông số từ datasheet của nhà sản xuất:\n"
               "  - Theo tài liệu kỹ thuật MPXV5010G, hàm truyền đặc tính ngõ ra lý thuyết (với nguồn nuôi Vs = 5.0 Vdc, nhiệt độ TA = 25°C) là: Vout = Vs * (0.09 * P + 0.04), trong đó P là áp suất đầu vào đo bằng kPa.\n"
               "  - Thực hiện biến đổi toán học tìm P từ điện áp ngõ ra Vout:\n"
               "    + Chia hai vế cho Vs: Vout / Vs = 0.09 * P + 0.04\n"
               "    + Trừ hai vế cho 0.04: (Vout / Vs) - 0.04 = 0.09 * P\n"
               "    + Chia hai vế cho 0.09: P = ((Vout / Vs) - 0.04) / 0.09 (kPa)\n"
               "  - Đồng nhất với mã nguồn: Biến `sensorVoltage` đại diện cho Vout. Giá trị số thực `5.0` đại diện cho nguồn cấp Vs = 5.0V. Biến `sensorVoltage` được khôi phục từ điện áp đọc tại chân ADC A0 của vi điều khiển thông qua hệ số nhân 1.5 (`sensorVoltage = voltage * 1.5`) nhằm bù đắp tỉ số chia áp của mạch cầu phân áp 10k/20k bên ngoài (giảm điện áp đi 1.5 lần, tức là V_adc = Vout * (20k / (10k + 20k)) = Vout * (2/3) => Vout = V_adc * 1.5) dùng để giới hạn ngõ ra cảm biến 5V không làm hỏng chân vào ADC 3.3V của vi điều khiển Arduino Nano 33 BLE.\n"
               "Như vậy, công thức trong mã nguồn là hoàn toàn chính xác, phản ánh trung thực giá trị áp lực vật lý đo được tại mặt nạ thở của bệnh nhân.\n"
               "Tương tự, đối với cảm biến lưu lượng kỹ thuật số Sensirion SFM3300-D, công thức chuyển đổi từ giá trị thô 14-bit đọc qua I2C (Raw_Flow) sang giá trị lưu lượng thực tế (slm) là:\n"
               "  `Flow (slm) = (Raw_Flow - 32768) / 120.0`\n"
               "Công thức này hoàn toàn chính xác theo datasheet của nhà sản xuất Sensirion với giá trị điểm không Offset Flow = 32768 và hệ số tỉ lệ Scale Factor Flow = 120 slm⁻¹.")
               
    add_bullet("Đo tốc độ quay của động cơ quạt thổi (RPM):\n"
               "Để kiểm soát quạt thổi khí WS4540 hoạt động chính xác và ổn định, xung phản hồi tốc độ (FG) của quạt được kết nối với chân D2 "
               "của Arduino sử dụng ngắt ngoài (External Interrupt) bắt cạnh sườn lên (RISING). Hàm ngắt countPulse() thực hiện đếm số lượng xung: \n"
               "  - pulseCount = pulseCount + 1\n"
               "Cứ sau mỗi chu kỳ 1000ms (1 giây), vi điều khiển sẽ tạm thời khóa ngắt, đọc giá trị biến đếm pulseCount, thiết lập lại pulseCount = 0 "
               "và kích hoạt lại ngắt. Vì quạt WS4540 phát ra 2 xung FG cho mỗi vòng quay, tốc độ động cơ (RPM - Revolutions Per Minute) được tính bằng công thức:\n"
               "  - RPM = (pulseCount / 2) * 60 (vòng/phút).")
               
    add_bullet("Cơ chế điều khiển tốc độ quạt (Blower Control):\n"
               "Tốc độ quạt được điều khiển bằng tín hiệu điều rộng xung (PWM) ghi ra chân D3 (analogWrite). Giá trị PWM từ 0 (dừng quạt) đến 255 (tốc độ tối đa). "
               "Trong đó, lệnh 'START' thiết lập giá trị PWM khởi động an toàn ban đầu là 55, lệnh 'STOP' đặt PWM về 0. "
               "Hệ thống cũng hỗ trợ nhận giá trị PWM tùy chọn từ thiết bị di động gửi xuống để tinh chỉnh dải áp suất đích.")
               
    add_bullet("Giao thức kết nối không dây BLE (Bluetooth Low Energy):\n"
               "Hệ thống vận hành như một thiết bị BLE ngoại vi (BLE Peripheral). Cấu hình chi tiết GATT Profile bao gồm:\n"
               "  - Tên thiết bị quảng bá: 'CPAP_VSSM'\n"
               "  - UUID của Dịch vụ CPAP: cb24858f-399f-4498-85e8-fea9d383d54f\n"
               "  - Đặc tính cảm biến (Sensor Characteristic): UUID 5e9e214b-124c-434d-84e5-018dccd35df1 (Hỗ trợ Read và Notify). Gửi định kỳ mỗi 200ms gói dữ liệu trạng thái chuẩn JSON: "
               '{"pressure": <giá trị>, "rpm": <giá trị>, "pwm": <giá trị>} để cập nhật trực tiếp lên ứng dụng di động.\n'
               "  - Đặc tính điều khiển (Action Characteristic): UUID 56debc28-acab-4184-8f86-1a9c887b220a (Hỗ trợ Write). Nhận các lệnh điều khiển chuỗi ký tự từ ứng dụng di động: "
               "'LED_ON' (bật LED chỉ thị), 'LED_OFF' (tắt LED chỉ thị), 'START' (chạy quạt), 'STOP' (tắt quạt), 'PWM:<giá trị>' (thiết lập tốc độ quạt trực tiếp).\n"
               "  - Tính năng an toàn: Khi mất kết nối BLE (disconnect), thiết bị tự động tắt quạt (PWM = 0) để bảo vệ an toàn cho bệnh nhân.")

    add_image_with_caption(doc, "image (2).png", "Hình 8: Thử nghiệm thực nghiệm đo đạc thông số và điều khiển tốc độ quạt thổi trên mô hình người giả kết nối với máy tính.", 4.0)

    # 12. Section 6: Technical and Economic Benefits
    add_heading_1("VI. HIỆU QUẢ ĐẠT ĐƯỢC VÀ ĐỊNH HƯỚNG PHÁT TRIỂN")
    p = doc.add_paragraph()
    p.add_run("1. Hiệu quả kỹ thuật và kinh tế:\n").bold = True
    add_bullet("Giá thành chế tạo cực kỳ thấp: Tổng chi phí linh kiện chỉ khoảng 7.095.000 VND, rẻ hơn từ 3 đến 8 lần so với các máy CPAP thương mại nhập khẩu trên thị trường.")
    add_bullet("Hệ thống tích hợp các cảm biến đo lường áp suất và lưu lượng khí thở thời gian thực, cung cấp dữ liệu hô hấp quan trọng cho bệnh nhân trong giấc ngủ.")
    add_bullet("Khả năng IoT hóa cao: Nhờ kết nối BLE và sự chuẩn bị của module ESP32, dữ liệu hô hấp có thể dễ dàng đồng bộ lên đám mây y tế SleepCare, cho phép bác sĩ chẩn đoán và theo dõi từ xa.")
    add_bullet("Thiết kế nhỏ gọn, vỏ mica cắt laser trong suốt giúp dễ dàng kiểm tra cơ khí trực quan, lắp đặt và sửa chữa nhanh chóng theo dạng module.")

    p = doc.add_paragraph()
    p.add_run("2. Lộ trình định hướng phát triển (Đưa vào nghiên cứu lâm sàng của VSSM):\n").bold = True
    add_bullet("Giai đoạn 1: Clinical Trial & Safety Certifications - Tiến hành thử nghiệm lâm sàng mẫu thử nghiệm để đánh giá độ an toàn y tế và hiệu chuẩn dải áp lực tối ưu dưới sự giám sát của GS.TS.KH. Dương Quý Sỹ.")
    add_bullet("Giai đoạn 2: Cloud Integration & Remote Monitoring - Hoàn thiện firmware cho vi điều khiển phụ ESP32 để kết nối Wi-Fi, đẩy dữ liệu trực tiếp lên cloud server mà không cần thông qua điện thoại trung gian.")
    add_bullet("Giai đoạn 3: AI-based Adaptation - Phát triển và nhúng thuật toán trí tuệ nhân tạo (AI/Machine Learning) trên ứng dụng di động hoặc máy chủ để tự động nhận dạng nhịp thở bất thường và đưa ra dự báo sớm cơn ngưng thở, từ đó điều chỉnh tự động áp suất quạt thổi (Auto-CPAP/APAP) thông minh hơn.")
    add_bullet("Giai đoạn 4: Sleep Ecosystem Integration - Tích hợp thiết bị vào hệ sinh thái chăm sóc giấc ngủ thông minh toàn diện cùng các cảm biến môi trường (nhiệt độ, ánh sáng, tiếng ồn phòng ngủ).")

    add_image_with_caption(doc, "image (5).png", "Hình 10: Sơ đồ lộ trình phát triển hệ thống SIPAP từ phiên bản thử nghiệm (MVP) đến hệ sinh thái chăm sóc giấc ngủ thông minh (Mô hình mô phỏng thiết kế dự kiến).", 6.0)

    # 13. Section 7: Patent Claims (Yêu cầu bảo hộ)
    add_heading_1("VII. YÊU CẦU BẢO HỘ")
    p = doc.add_paragraph()
    p.add_run("1. Hệ thống hỗ trợ thở áp lực dương liên tục (CPAP) tương tác thông minh (SIPAP) dùng trong hỗ trợ điều trị hội chứng ngưng thở khi ngủ, hệ thống bao gồm:")
    add_bullet("Một khối tạo áp lực dương gồm quạt thổi không chổi than 12VDC và bộ điều khiển tốc độ bằng xung PWM;")
    add_bullet("Một khối cảm biến đo lường tích hợp gồm cảm biến áp suất analog kết nối qua cầu phân áp để giám sát áp suất mặt nạ thở, và cảm biến lưu lượng dòng khí giao tiếp chuẩn I2C;")
    add_bullet("Một khối vi điều khiển trung tâm tích hợp module truyền thông Bluetooth Low Energy (BLE), được kết nối điện với khối tạo áp lực và khối cảm biến;")
    add_bullet("Đặc trưng ở chỗ: Vi điều khiển trung tâm được lập trình để liên tục đọc tín hiệu áp suất mặt nạ thở từ cảm biến áp suất, tính toán lưu lượng từ cảm biến dòng khí, đồng thời đo tốc độ quay RPM của quạt thổi thông qua đếm số lượng xung phản hồi tốc độ quay nối vào chân ngắt ngoài của vi điều khiển, và định kỳ truyền chuỗi dữ liệu định dạng JSON chứa các thông số áp suất, RPM và giá trị PWM quạt thổi qua giao tiếp BLE GATT Profile dưới tên quảng bá 'SIPAP' để giám sát thời gian thực.")
    
    p = doc.add_paragraph()
    p.add_run("2. Hệ thống hỗ trợ thở theo điểm yêu cầu bảo hộ 1, đặc trưng ở chỗ: ")
    p.add_run("Thiết kế phần cứng tích hợp sẵn cổng kết nối chờ và thiết lập mã nguồn dự phòng để tích hợp cảm biến đo nồng độ bão hòa oxy máu SpO2 quang học ở phiên bản nâng cấp.")
    
    p = doc.add_paragraph()
    p.add_run("3. Hệ thống hỗ trợ thở theo điểm yêu cầu bảo hộ 1 hoặc 2, đặc trưng ở chỗ: ")
    p.add_run("Vi điều khiển trung tâm được cấu hình tính năng bảo vệ an toàn tự động ngắt quạt thổi (đặt PWM = 0) ngay lập tức khi phát hiện mất kết nối không dây BLE với ứng dụng di động.")
    
    p = doc.add_paragraph()
    p.add_run("4. Hệ thống hỗ trợ thở theo điểm yêu cầu bảo hộ 1, đặc trưng ở chỗ: ")
    p.add_run("Vỏ máy thở và buồng chứa quạt tạo áp lực dương được chế tạo nguyên khối từ các tấm mica trong suốt cắt laser có độ dày từ 3-5mm để đảm bảo độ kín khí và khả năng quan sát trực quan các thành phần phần cứng bên trong.")
    
    p = doc.add_paragraph()
    p.add_run("5. Hệ thống hỗ trợ thở theo điểm yêu cầu bảo hộ 1, đặc trưng ở chỗ: ")
    p.add_run("Bảo hộ tên gọi độc quyền sáng chế 'SIPAP' (Smart Interactive PAP - Thiết bị hỗ trợ thở áp lực dương tương tác thông minh) do GS.TS.KH. Dương Quý Sỹ đề xuất sáng tạo gắn liền với hệ sinh thái giải pháp kỹ thuật y tế nêu trên.")

    # Save to file
    out_dir = r"c:\Users\ADMIN\OneDrive\Máy tính\Master_2024\Cpap\document"
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, "v1.docx")
    doc.save(out_path)
    print(f"Document saved successfully at {out_path}!")

if __name__ == "__main__":
    create_document()
