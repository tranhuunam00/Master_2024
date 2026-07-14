import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def clean_latex_math(text):
    # Convert \text{...} to ...
    text = re.sub(r'\\text\{([^}]+)\}', r'\1', text)
    # Convert \frac{A}{B} to A / B
    text = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'\1 / \2', text)
    
    # Strip any \mathrm, \mathit, \mathbf, \mathsf
    text = re.sub(r'\\math[a-z]+\{([^}]+)\}', r'\1', text)
    
    # Convert subscripts like _{A} to (A)
    text = re.sub(r'_\{([^}]+)\}', r' (\1)', text)
    
    # Remove $ signs
    text = text.replace('$', '')
    
    # Convert operators to professional Unicode math characters
    text = text.replace(r'\le', ' ≤ ')
    text = text.replace(r'\ge', ' ≥ ')
    text = text.replace(r'\implies', ' ⇒ ')
    text = text.replace(r'\times', ' × ')
    text = text.replace(r'\approx', ' ≈ ')
    text = text.replace(r'\in', ' ∈ ')
    
    # Convert sum limits to friendly text
    text = text.replace(r'\sum_{i=1}^{8}', 'Tổng (i = 1 đến 8) ')
    text = text.replace(r'\sum_{i=1}^{17}', 'Tổng (i = 1 đến 17) ')
    text = text.replace(r'\sum_{j=1}^{7}', 'Tổng (j = 1 đến 7) ')
    text = text.replace(r'\sum_{i ∈ Thuận}', 'Tổng (i ∈ Thuận) ')
    text = text.replace(r'\sum_{i ∈ Ngược}', 'Tổng (i ∈ Ngược) ')
    text = text.replace(r'\sum_{i \in Thuận}', 'Tổng (i ∈ Thuận) ')
    text = text.replace(r'\sum_{i \in Ngược}', 'Tổng (i ∈ Ngược) ')
    text = text.replace(r'\sum', 'Tổng ')
    
    # Cleanup subscripts and variables
    text = text.replace(r'_{STOP-BANG}', ' STOP-BANG')
    text = text.replace(r'_{ESS}', ' ESS')
    text = text.replace(r'_{Zung}', ' Zung')
    text = text.replace(r'_{Hamilton}', ' Hamilton')
    text = text.replace(r'_{PSQI}', ' PSQI')
    text = text.replace(r'_{ISI}', ' ISI')
    text = text.replace(r'_{Pichot-Fatigue}', ' Pichot-Fatigue')
    text = text.replace(r'_{Pichot-QD}', ' Pichot-QD')
    text = text.replace(r'_{GDS-15}', ' GDS-15')
    text = text.replace(r'_{Age}', ' Age')
    text = text.replace(r'_{Neck}', ' Neck')
    text = text.replace(r'_{BMI}', ' BMI')
    text = text.replace(r'_i', ' i')
    text = text.replace(r'_j', ' j')
    
    # Cleanup others
    text = text.replace(r'\quad', ' ')
    text = text.replace(r'^2', '²')
    text = text.replace(r'\mid', ' | ')
    text = text.replace(r'\ll', ' << ')
    
    # Layout and font size overrides cleanups
    text = text.replace(r'\\', '')
    text = re.sub(r'\\vspace\{[^}]+\}', '', text)
    text = text.replace(r'\noindent', '')
    text = text.replace(r'\large', '')
    text = text.replace(r'\Large', '')
    text = text.replace(r'\small', '')
    text = text.replace(r'\normalsize', '')
    text = text.replace(r'\begin{center}', '')
    text = text.replace(r'\end{center}', '')
    
    # Double spaces cleanup
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def clean_latex_markup(text):
    text = clean_latex_math(text)
    # Remove simple markup but keep the content
    text = re.sub(r'\\textbf\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\textcolor\{[a-zA-Z_]+\}\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\text[a-zA-Z]+\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\href\{[^}]+\}\{([^}]+)\}', r'\1', text)
    text = text.replace(r'\noindent', '')
    text = text.replace(r'~', ' ')
    return text.strip()

def add_paragraph_with_runs(p, text):
    text = clean_latex_math(text)
    # Simple parsing of \textbf{} and \textit{} to add runs
    pattern = re.compile(r'(\\textbf\{[^}]+\}|\\textit\{[^}]+\}|\\texttt\{[^}]+\}|[^\\]+)')
    pos = 0
    while pos < len(text):
        match = re.search(r'\\(textbf|textit|texttt|textcolor)\{([^}]+)\}', text[pos:])
        if match:
            start = pos + match.start()
            end = pos + match.end()
            # Add text before match
            if start > pos:
                p.add_run(text[pos:start])
            
            # Add formatted run
            fmt = match.group(1)
            content = match.group(2)
            if fmt == 'textcolor':
                # Color match is actually \textcolor{colorname}{content}
                # Let's extract content and color
                color_match = re.search(r'\\textcolor\{([a-zA-Z]+)\}\{([^}]+)\}', text[pos:])
                if color_match:
                    color_name = color_match.group(1)
                    content = color_match.group(2)
                    run = p.add_run(content)
                    if color_name.lower() in ['navy', 'blue']:
                        run.font.color.rgb = RGBColor(0, 51, 102)
                    elif color_name.lower() == 'red':
                        run.font.color.rgb = RGBColor(204, 0, 0)
                    end = pos + color_match.end()
            else:
                run = p.add_run(content)
                if fmt == 'textbf':
                    run.bold = True
                elif fmt == 'textit':
                    run.italic = True
                elif fmt == 'texttt':
                    run.font.name = 'Courier New'
            pos = end
        else:
            p.add_run(text[pos:])
            break
            
    # Set default font for all runs in the paragraph
    for run in p.runs:
        if run.font.name != 'Courier New':
            run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def compile_tex_to_docx(tex_path, docx_path):
    print("Reading LaTeX file...")
    with open(tex_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    
    # Configure page setup (margins: 3.0cm Left, 2.0cm others)
    for section in doc.sections:
        section.top_margin = Inches(0.79)     # 2.0 cm
        section.bottom_margin = Inches(0.79)  # 2.0 cm
        section.left_margin = Inches(1.18)    # 3.0 cm
        section.right_margin = Inches(0.79)   # 2.0 cm

    # Configure base style
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    # Split lines
    lines = content.split('\n')
    
    # State tracking
    in_document = False
    in_itemize = False
    in_enumerate = False
    in_equation = False
    in_figure = False
    in_center = False
    list_counter = 0
    figure_img = ""
    figure_caption = ""

    # Parse title info
    title = "GIẢI PHÁP HỮU ÍCH"

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip comments
        if line.startswith('%'):
            i += 1
            continue

        if r'\begin{document}' in line:
            in_document = True
            i += 1
            continue
            
        if not in_document:
            i += 1
            continue

        if r'\end{document}' in line:
            break

        # Check section
        section_match = re.match(r'\\section\{\\textcolor\{navy\}\{([^}]+)\}\}', line)
        if not section_match:
            section_match = re.match(r'\\section\{([^}]+)\}', line)
        if section_match:
            title_text = section_match.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(title_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102) # Navy color
            i += 1
            continue

        # Check subsection
        sub_match = re.match(r'\\subsection\{([^}]+)\}', line)
        if sub_match:
            sub_text = sub_match.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(sub_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 102, 204) # Lighter blue
            i += 1
            continue

        # Check subsubsection
        subsub_match = re.match(r'\\subsubsection\*?\{([^}]+)\}', line)
        if subsub_match:
            subsub_text = subsub_match.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_latex_markup(subsub_text))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            i += 1
            continue

        # Lists environment
        if r'\begin{itemize}' in line:
            in_itemize = True
            i += 1
            continue
        if r'\end{itemize}' in line:
            in_itemize = False
            i += 1
            continue
        if r'\begin{enumerate}' in line:
            in_enumerate = True
            list_counter = 0
            i += 1
            continue
        if r'\end{enumerate}' in line:
            in_enumerate = False
            i += 1
            continue

        # Check item
        if line.startswith(r'\item'):
            content_text = line.replace(r'\item', '').strip()
            # If item is multi-line, collect the lines
            while i + 1 < len(lines) and not lines[i+1].strip().startswith(r'\item') and not lines[i+1].strip().startswith(r'\end{') and not lines[i+1].strip().startswith(r'\begin{'):
                i += 1
                content_text += " " + lines[i].strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            if in_enumerate:
                list_counter += 1
                p.paragraph_format.left_indent = Inches(0.4)
                # First run is number
                run_num = p.add_run(f"{list_counter}. ")
                run_num.font.name = 'Times New Roman'
                run_num.font.size = Pt(12)
                run_num.bold = True
            else:
                p.paragraph_format.left_indent = Inches(0.4)
                run_bullet = p.add_run("• ")
                run_bullet.font.name = 'Times New Roman'
                run_bullet.font.size = Pt(12)
            
            add_paragraph_with_runs(p, content_text)
            i += 1
            continue

        # Equation environment
        if r'\begin{equation}' in line:
            in_equation = True
            equation_content = ""
            i += 1
            continue
        if r'\end{equation}' in line:
            in_equation = False
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(clean_latex_markup(equation_content))
            run.font.name = 'Consolas'
            run.font.size = Pt(11)
            run.italic = True
            i += 1
            continue
        if in_equation:
            equation_content += " " + line
            i += 1
            continue

        # Figure environment
        if r'\begin{figure}' in line:
            in_figure = True
            figure_img = ""
            figure_caption = ""
            i += 1
            continue
        if r'\end{figure}' in line:
            in_figure = False
            # Insert centered image
            if figure_img:
                img_path = os.path.join(os.path.dirname(tex_path), figure_img)
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(12)
                    p_img.paragraph_format.space_after = Pt(6)
                    run_img = p_img.add_run()
                    # Scale image reasonably
                    if 'logo' in figure_img:
                        run_img.add_picture(img_path, width=Inches(1.8))
                    else:
                        run_img.add_picture(img_path, width=Inches(5.0))
                else:
                    print(f"Warning: Image file not found: {img_path}")
            
            # Insert caption
            if figure_caption:
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(12)
                run_cap = p_cap.add_run(clean_latex_markup(figure_caption))
                run_cap.font.name = 'Times New Roman'
                run_cap.font.size = Pt(10.5)
                run_cap.italic = True
            
            i += 1
            continue

        if in_figure:
            if r'\includegraphics' in line:
                match = re.search(r'\\includegraphics\[[^\]]*\]\{([^}]+)\}', line)
                if not match:
                    match = re.search(r'\\includegraphics\{([^}]+)\}', line)
                if match:
                    figure_img = match.group(1)
            elif r'\caption' in line:
                match = re.search(r'\\caption\{([^}]+)\}', line)
                if match:
                    figure_caption = match.group(1)
            i += 1
            continue

        # Normal text paragraph
        if line != "":
            # Check if it is center environment tag
            if r'\begin{center}' in line:
                in_center = True
                i += 1
                continue
            if r'\end{center}' in line:
                in_center = False
                i += 1
                continue
                
            # Collect full paragraph lines (until empty line or LaTeX command)
            para_text = line
            while i + 1 < len(lines) and lines[i+1].strip() != "" and not lines[i+1].strip().startswith(r'\section') and not lines[i+1].strip().startswith(r'\subsection') and not lines[i+1].strip().startswith(r'\begin{') and not lines[i+1].strip().startswith(r'\item') and not lines[i+1].strip().startswith(r'\end{') and not lines[i+1].strip().startswith(r'\noindent'):
                i += 1
                para_text += " " + lines[i].strip()
            
            # Split by \\ to handle line breaks
            parts = para_text.split(r'\\')
            for part in parts:
                part = part.strip()
                if part == "":
                    continue
                # Skip if it contains center tags
                if r'\begin{center}' in part:
                    in_center = True
                    continue
                if r'\end{center}' in part:
                    in_center = False
                    continue
                
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                if in_center:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                add_paragraph_with_runs(p, part)

        i += 1

    print("Saving DOCX file...")
    try:
        doc.save(docx_path)
        print("Compilation completed successfully!")
    except PermissionError:
        counter = 1
        while True:
            suffix = f"_{counter}" if counter > 1 else "_new"
            fallback_name = f"giai_phap_huu_ich_SleepCare_v1{suffix}.docx"
            fallback_path = os.path.join(os.path.dirname(docx_path), fallback_name)
            try:
                doc.save(fallback_path)
                print(f"Permission denied. Saved successfully to: {fallback_name}")
                print("Compilation completed successfully!")
                break
            except PermissionError:
                counter += 1
                if counter > 50:
                    print("ERROR: All fallback filenames are locked. Please close Word and retry!")
                    raise

if __name__ == '__main__':
    current_dir = os.path.dirname(os.path.abspath(__file__))
    tex_file = os.path.join(current_dir, "giai_phap_huu_ich_SleepCare_v1.tex")
    docx_file = os.path.join(current_dir, "giải phap hữu ích SleepCare v1.docx")
    compile_tex_to_docx(tex_file, docx_file)
