import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import latex2mathml.converter
from lxml import etree

# Load XSLT for MathML → OMML conversion (one-time)
XSLT_PATH = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"
_xslt_tree = etree.parse(XSLT_PATH)
_xslt_transform = etree.XSLT(_xslt_tree)

def latex_to_omml(latex_str):
    """Convert a LaTeX math string to an OMML XML element for Word with post-processing."""
    try:
        mathml_str = latex2mathml.converter.convert(latex_str)
        mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
        omml_tree = _xslt_transform(mathml_tree)
        omml_root = omml_tree.getroot()
        
        # Post-process OMML to fix empty tags (avoid dotted square placeholders) and force under-over summation limits
        namespaces = {'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'}
        
        # 1. Force summation limit locations to be under/over (undOvr)
        for limloc in omml_root.xpath('//m:limLoc', namespaces=namespaces):
            limloc.set('{http://schemas.openxmlformats.org/officeDocument/2006/math}val', 'undOvr')
            
        # 2. Sibling swallowing logic: put the summation operands inside <m:e> to prevent empty operand dotted squares
        for parent in omml_root.xpath('//m:nary/..', namespaces=namespaces):
            children = list(parent)
            i = 0
            while i < len(children):
                child = children[i]
                if child.tag == '{http://schemas.openxmlformats.org/officeDocument/2006/math}nary':
                    nary = child
                    e_list = nary.xpath('m:e', namespaces=namespaces)
                    if e_list:
                        e_elem = e_list[0]
                        # Swallow succeeding siblings (up to next summation or binary operators)
                        j = i + 1
                        swallowed = []
                        while j < len(children):
                            next_sibling = children[j]
                            if next_sibling.tag == '{http://schemas.openxmlformats.org/officeDocument/2006/math}nary':
                                break
                            t_text = next_sibling.xpath('m:t', namespaces=namespaces)
                            if t_text and t_text[0].text in ['=', ' + ', ' - ']:
                                break
                            swallowed.append(next_sibling)
                            j += 1
                        
                        for elem in swallowed:
                            parent.remove(elem)
                            e_elem.append(elem)
                        
                        children = list(parent)
                i += 1
                
        # 3. Clean empty sup/sub tags to avoid empty dotted squares
        for nary in omml_root.xpath('//m:nary', namespaces=namespaces):
            sup = nary.xpath('m:sup', namespaces=namespaces)
            if sup and not (len(sup[0]) > 0 or sup[0].text):
                suphide = nary.xpath('m:naryPr/m:supHide', namespaces=namespaces)
                if suphide:
                    suphide[0].set('{http://schemas.openxmlformats.org/officeDocument/2006/math}val', 'on')
                nary.remove(sup[0])
                
            sub = nary.xpath('m:sub', namespaces=namespaces)
            if sub and not (len(sub[0]) > 0 or sub[0].text):
                subhide = nary.xpath('m:naryPr/m:subHide', namespaces=namespaces)
                if subhide:
                    subhide[0].set('{http://schemas.openxmlformats.org/officeDocument/2006/math}val', 'on')
                nary.remove(sub[0])
                
        return omml_root
    except Exception as e:
        print(f"Warning: Could not convert equation to OMML: {e}")
        return None

def add_equation_to_doc(doc, latex_str):
    """Add a centered Word equation paragraph from LaTeX source."""
    omml_element = latex_to_omml(latex_str)
    if omml_element is not None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p._element.append(omml_element)
        return True
    return False

def create_element(name):
    return OxmlElement(name)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def clean_latex_math(text):
    # Convert \text{...} to ...
    text = re.sub(r'\\text\{([^}]+)\}', r'\1', text)
    # Convert \frac{A}{B} to A / B
    text = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'\1 / \2', text)
    
    # Strip any \mathrm, \mathit, \mathbf, \mathsf
    text = re.sub(r'\\math[a-z]+\{([^}]+)\}', r'\1', text)
    
    # Remove $ signs
    text = text.replace('$', '')
    
    # === STEP 1: Process \sum FIRST (before subscript/superscript conversion) ===
    # \sum_{lower}^{upper} → Σ(lower→upper)
    text = re.sub(r'\\sum_\{([^}]+)\}\^\{([^}]+)\}', r'Σ(\1→\2)', text)
    # \sum_{lower} → Σ(lower)
    text = re.sub(r'\\sum_\{([^}]+)\}', r'Σ(\1)', text)
    # Bare \sum → Σ
    text = text.replace(r'\sum', 'Σ')
    
    # === STEP 2: Convert escaped braces \{ \} → { } ===
    text = text.replace(r'\{', '{')
    text = text.replace(r'\}', '}')
    
    # === STEP 3: Convert superscripts ^{n} → ⁿ ===
    superscript_map = {'0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
                       '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹'}
    def replace_superscript(m):
        content = m.group(1)
        result = ''
        for ch in content:
            result += superscript_map.get(ch, ch)
        return result
    text = re.sub(r'\^\{([^}]+)\}', replace_superscript, text)
    text = text.replace(r'^2', '²')
    
    # === STEP 4: Convert subscripts _{A} → (A) for multi-char, or Unicode for single-char ===
    subscript_map = {
        '0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄',
        '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉',
        'a': 'ₐ', 'e': 'ₑ', 'h': 'ₕ', 'i': 'ᵢ', 'j': 'ⱼ',
        'k': 'ₖ', 'l': 'ₗ', 'm': 'ₘ', 'n': 'ₙ', 'o': 'ₒ',
        'p': 'ₚ', 'r': 'ᵣ', 's': 'ₛ', 't': 'ₜ', 'u': 'ᵤ',
        'v': 'ᵥ', 'x': 'ₓ',
    }
    def replace_subscript_braces(m):
        content = m.group(1)
        # If all chars have Unicode subscript equivalents, use them
        converted = ''
        for ch in content:
            if ch in subscript_map:
                converted += subscript_map[ch]
            else:
                # Fallback to parenthesized notation for complex subscripts
                return f'({content})'
        return converted
    text = re.sub(r'_\{([^}]+)\}', replace_subscript_braces, text)
    
    # Single-char subscripts: _i → ᵢ, _j → ⱼ
    text = re.sub(r'_([a-z0-9])', lambda m: subscript_map.get(m.group(1), m.group(1)), text)
    
    # === STEP 5: Convert operators to Unicode ===
    text = text.replace(r'\le', ' ≤ ')
    text = text.replace(r'\ge', ' ≥ ')
    text = text.replace(r'\implies', ' ⇒ ')
    text = text.replace(r'\times', ' × ')
    text = text.replace(r'\approx', ' ≈ ')
    text = text.replace(r'\in', ' ∈ ')
    
    # === STEP 6: Cleanup others ===
    text = text.replace(r'\quad', ' ')
    text = text.replace(r'\mid', ' | ')
    text = text.replace(r'\ll', ' << ')
    
    # Layout and font size overrides cleanups
    text = text.replace(r'\\', '')
    text = re.sub(r'\\vspace\{[^}]+\}', '', text)
    text = text.replace(r'\noindent', '')
    text = text.replace(r'\large', '')
    text = text.replace(r'\Large', '')
    text = text.replace(r'\small', '')
    text = text.replace(r'\normalsize', '')
    text = text.replace(r'\begin{center}', '')
    text = text.replace(r'\end{center}', '')
    
    # Double spaces cleanup
    text = re.sub(r'\s+', ' ', text)
    return text.strip()

def clean_latex_markup(text):
    text = clean_latex_math(text)
    # Remove simple markup but keep the content
    text = re.sub(r'\\textbf\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\texttt\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\textcolor\{[a-zA-Z_]+\}\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\text[a-zA-Z]+\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\href\{[^}]+\}\{([^}]+)\}', r'\1', text)
    text = text.replace(r'\noindent', '')
    text = text.replace(r'~', ' ')
    return text.strip()

def add_paragraph_with_runs(p, text):
    text = clean_latex_math(text)
    # Simple parsing of \textbf{} and \textit{} to add runs
    pattern = re.compile(r'(\\textbf\{[^}]+\}|\\textit\{[^}]+\}|\\texttt\{[^}]+\}|[^\\]+)')
    pos = 0
    while pos < len(text):
        match = re.search(r'\\(textbf|textit|texttt|textcolor)\{([^}]+)\}', text[pos:])
        if match:
            start = pos + match.start()
            end = pos + match.end()
            # Add text before match
            if start > pos:
                p.add_run(text[pos:start])
            
            # Add formatted run
            fmt = match.group(1)
            content = match.group(2)
            if fmt == 'textcolor':
                # Color match is actually \textcolor{colorname}{content}
                # Let's extract content and color
                color_match = re.search(r'\\textcolor\{([a-zA-Z]+)\}\{([^}]+)\}', text[pos:])
                if color_match:
                    color_name = color_match.group(1)
                    content = color_match.group(2)
                    run = p.add_run(content)
                    if color_name.lower() in ['navy', 'blue']:
                        run.font.color.rgb = RGBColor(0, 51, 102)
                    elif color_name.lower() == 'red':
                        run.font.color.rgb = RGBColor(204, 0, 0)
                    end = pos + color_match.end()
            else:
                run = p.add_run(content)
                if fmt == 'textbf':
                    run.bold = True
                elif fmt == 'textit':
                    run.italic = True
                elif fmt == 'texttt':
                    run.font.name = 'Courier New'
            pos = end
        else:
            p.add_run(text[pos:])
            break
            
    # Set default font for all runs in the paragraph
    for run in p.runs:
        if run.font.name != 'Courier New':
            run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def compile_tex_to_docx(tex_path, docx_path):
    print("Reading LaTeX file...")
    with open(tex_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()
    
    # Configure page setup (margins: 3.0cm Left, 2.0cm others)
    for section in doc.sections:
        section.top_margin = Inches(0.79)     # 2.0 cm
        section.bottom_margin = Inches(0.79)  # 2.0 cm
        section.left_margin = Inches(1.18)    # 3.0 cm
        section.right_margin = Inches(0.79)   # 2.0 cm

    # Configure base style
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    # Split lines
    lines = content.split('\n')
    
    # State tracking
    in_document = False
    in_itemize = False
    in_enumerate = False
    in_equation = False
    in_figure = False
    in_center = False
    list_counter = 0
    figure_img = ""
    figure_caption = ""

    # Parse title info
    title = "GIẢI PHÁP HỮU ÍCH"

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip comments
        if line.startswith('%'):
            i += 1
            continue

        if r'\begin{document}' in line:
            in_document = True
            i += 1
            continue
            
        if not in_document:
            i += 1
            continue

        if r'\end{document}' in line:
            break

        # Check section
        section_match = re.match(r'\\section\{\\textcolor\{navy\}\{([^}]+)\}\}', line)
        if not section_match:
            section_match = re.match(r'\\section\{([^}]+)\}', line)
        if section_match:
            title_text = section_match.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(title_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102) # Navy color
            i += 1
            continue

        # Check subsection
        sub_match = re.match(r'\\subsection\{([^}]+)\}', line)
        if sub_match:
            sub_text = sub_match.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(sub_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 102, 204) # Lighter blue
            i += 1
            continue

        # Check subsubsection
        subsub_match = re.match(r'\\subsubsection\*?\{([^}]+)\}', line)
        if subsub_match:
            subsub_text = subsub_match.group(1)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_latex_markup(subsub_text))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
            i += 1
            continue

        # Lists environment
        if r'\begin{itemize}' in line:
            in_itemize = True
            i += 1
            continue
        if r'\end{itemize}' in line:
            in_itemize = False
            i += 1
            continue
        if r'\begin{enumerate}' in line:
            in_enumerate = True
            list_counter = 0
            i += 1
            continue
        if r'\end{enumerate}' in line:
            in_enumerate = False
            i += 1
            continue

        # Check item
        if line.startswith(r'\item'):
            content_text = line.replace(r'\item', '').strip()
            # If item is multi-line, collect the lines
            while i + 1 < len(lines) and not lines[i+1].strip().startswith(r'\item') and not lines[i+1].strip().startswith(r'\end{') and not lines[i+1].strip().startswith(r'\begin{'):
                i += 1
                content_text += " " + lines[i].strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            if in_enumerate:
                list_counter += 1
                p.paragraph_format.left_indent = Inches(0.4)
                # First run is number
                run_num = p.add_run(f"{list_counter}. ")
                run_num.font.name = 'Times New Roman'
                run_num.font.size = Pt(12)
                run_num.bold = True
            else:
                p.paragraph_format.left_indent = Inches(0.4)
                run_bullet = p.add_run("• ")
                run_bullet.font.name = 'Times New Roman'
                run_bullet.font.size = Pt(12)
            
            add_paragraph_with_runs(p, content_text)
            i += 1
            continue

        # Equation environment
        if r'\begin{equation}' in line:
            in_equation = True
            equation_content = ""
            i += 1
            continue
        if r'\end{equation}' in line:
            in_equation = False
            # Try OMML equation first, fallback to plain text
            if not add_equation_to_doc(doc, equation_content.strip()):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(clean_latex_markup(equation_content))
                run.font.name = 'Consolas'
                run.font.size = Pt(11)
                run.italic = True
            i += 1
            continue
        if in_equation:
            equation_content += " " + line
            i += 1
            continue

        # Figure environment
        if r'\begin{figure}' in line:
            in_figure = True
            figure_img = ""
            figure_caption = ""
            i += 1
            continue
        if r'\end{figure}' in line:
            in_figure = False
            # Insert centered image
            if figure_img:
                img_path = os.path.join(os.path.dirname(tex_path), figure_img)
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(12)
                    p_img.paragraph_format.space_after = Pt(6)
                    run_img = p_img.add_run()
                    # Scale image reasonably
                    if 'logo' in figure_img:
                        run_img.add_picture(img_path, width=Inches(1.8))
                    else:
                        run_img.add_picture(img_path, width=Inches(5.0))
                else:
                    print(f"Warning: Image file not found: {img_path}")
            
            # Insert caption
            if figure_caption:
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_after = Pt(12)
                run_cap = p_cap.add_run(clean_latex_markup(figure_caption))
                run_cap.font.name = 'Times New Roman'
                run_cap.font.size = Pt(10.5)
                run_cap.italic = True
            
            i += 1
            continue

        if in_figure:
            if r'\includegraphics' in line:
                match = re.search(r'\\includegraphics\[[^\]]*\]\{([^}]+)\}', line)
                if not match:
                    match = re.search(r'\\includegraphics\{([^}]+)\}', line)
                if match:
                    figure_img = match.group(1)
            elif r'\caption' in line:
                match = re.search(r'\\caption\{([^}]+)\}', line)
                if match:
                    figure_caption = match.group(1)
            i += 1
            continue

        # Normal text paragraph
        if line != "":
            # Check if it is center environment tag
            if r'\begin{center}' in line:
                in_center = True
                i += 1
                continue
            if r'\end{center}' in line:
                in_center = False
                i += 1
                continue
                
            # Collect full paragraph lines (until empty line or LaTeX command)
            para_text = line
            while i + 1 < len(lines) and lines[i+1].strip() != "" and not lines[i+1].strip().startswith(r'\section') and not lines[i+1].strip().startswith(r'\subsection') and not lines[i+1].strip().startswith(r'\begin{') and not lines[i+1].strip().startswith(r'\item') and not lines[i+1].strip().startswith(r'\end{') and not lines[i+1].strip().startswith(r'\noindent'):
                i += 1
                para_text += " " + lines[i].strip()
            
            # Split by \\ to handle line breaks
            parts = para_text.split(r'\\')
            for part in parts:
                part = part.strip()
                if part == "":
                    continue
                # Skip if it contains center tags
                if r'\begin{center}' in part:
                    in_center = True
                    continue
                if r'\end{center}' in part:
                    in_center = False
                    continue
                
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                if in_center:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                add_paragraph_with_runs(p, part)

        i += 1

    print("Saving DOCX file...")
    try:
        doc.save(docx_path)
        print("Compilation completed successfully!")
    except PermissionError:
        counter = 1
        while True:
            suffix = f"_{counter}" if counter > 1 else "_new"
            fallback_name = f"giai_phap_huu_ich_SleepCare_v1{suffix}.docx"
            fallback_path = os.path.join(os.path.dirname(docx_path), fallback_name)
            try:
                doc.save(fallback_path)
                print(f"Permission denied. Saved successfully to: {fallback_name}")
                print("Compilation completed successfully!")
                break
            except PermissionError:
                counter += 1
                if counter > 50:
                    print("ERROR: All fallback filenames are locked. Please close Word and retry!")
                    raise

if __name__ == '__main__':
    current_dir = os.path.dirname(os.path.abspath(__file__))
    tex_file = os.path.join(current_dir, "giai_phap_huu_ich_SleepCare_v1.tex")
    docx_file = os.path.join(current_dir, "giải phap hữu ích SleepCare v1.docx")
    compile_tex_to_docx(tex_file, docx_file)
