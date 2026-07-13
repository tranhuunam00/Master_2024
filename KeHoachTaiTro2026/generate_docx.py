import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = Document()
    
    # 1. Page Margins Setup (Standard Vietnamese administrative layout: Left 3cm, Right 2cm, Top 2cm, Bottom 2cm)
    # 1 inch = 2.54 cm -> 3cm = 1.18 in, 2cm = 0.79 in
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.79)
        
    # 2. Typography and Styles Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    # Custom Paragraph format for Normal style
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # 3. Header Administrative Block (using Table for neat side-by-side alignment)
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    
    # Set widths for columns
    header_table.columns[0].width = Inches(2.8)
    header_table.columns[1].width = Inches(3.6)
    
    cell_left = header_table.cell(0, 0)
    cell_right = header_table.cell(0, 1)
    
    # Left Header (Issuing Authority)
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_left.paragraph_format.space_after = Pt(2)
    p_left.paragraph_format.line_spacing = 1.1
    
    run_dept = p_left.add_run("BỘ KHOA HỌC VÀ CÔNG NGHỆ\n")
    run_dept.font.size = Pt(12)
    
    run_fund = p_left.add_run("QUỸ PHÁT TRIỂN KHOA HỌC VÀ\nCÔNG NGHỆ QUỐC GIA\n")
    run_fund.bold = True
    run_fund.font.size = Pt(12)
    
    run_line_left = p_left.add_run("———————\n")
    run_line_left.font.size = Pt(10)
    
    run_no = p_left.add_run("Số:       /TB-QPTKHCNQG")
    run_no.font.size = Pt(12)
    
    # Right Header (National Title & Date)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_right.paragraph_format.space_after = Pt(2)
    p_right.paragraph_format.line_spacing = 1.1
    
    run_nation1 = p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    run_nation1.bold = True
    run_nation1.font.size = Pt(12)
    
    run_nation2 = p_right.add_run("Độc lập - Tự do - Hạnh phúc\n")
    run_nation2.bold = True
    run_nation2.font.size = Pt(13)
    
    run_line_right = p_right.add_run("___________________\n\n")
    run_line_right.font.size = Pt(10)
    
    run_date = p_right.add_run("Hà Nội, ngày 08 tháng 07 năm 2026")
    run_date.italic = True
    run_date.font.size = Pt(13)
    
    # Empty space after header block
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # 4. Title of the Announcement
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(12)
    p_title.paragraph_format.line_spacing = 1.15
    
    run_title_type = p_title.add_run("THÔNG BÁO\n")
    run_title_type.bold = True
    run_title_type.font.size = Pt(14)
    
    run_title_main = p_title.add_run("Kế hoạch tài trợ nhiệm vụ nghiên cứu phát triển công nghệ năm 2026")
    run_title_main.bold = True
    run_title_main.font.size = Pt(14)
    
    # 5. Citations (Căn cứ)
    citations = [
        "Căn cứ Nghị quyết số 57-NQ/TW ngày 22 tháng 12 năm 2024 của Bộ Chính trị về đột phá phát triển khoa học, công nghệ, đổi mới sáng tạo và chuyển đổi số quốc gia;",
        "Căn cứ Luật Khoa học, công nghệ và đổi mới sáng tạo số 93/2025/QH15;",
        "Căn cứ Luật Ngân sách nhà nước số 89/2025/QH15;",
        "Căn cứ Nghị định số 229/2026/NĐ-CP ngày 25 tháng 6 năm 2026 của Chính phủ quy định về tổ chức và hoạt động của Quỹ Phát triển khoa học và công nghệ Quốc gia;",
        "Căn cứ Nghị định số 265/2025/NĐ-CP ngày 14 tháng 10 năm 2025 của Chính phủ quy định chi tiết và hướng dẫn thi hành một số điều của Luật Khoa học, công nghệ và đổi mới sáng tạo về tài chính và đầu tư trong khoa học, công nghệ và đổi mới sáng tạo;",
        "Căn cứ Nghị định số 267/2025/NĐ-CP ngày 14 tháng 10 năm 2025 của Chính phủ quy định chi tiết và hướng dẫn một số điều của Luật Khoa học, công nghệ và đổi mới sáng tạo về chương trình, nhiệm vụ khoa học, công nghệ và đổi mới sáng tạo và một số quy định về thúc đẩy hoạt động nghiên cứu khoa học, phát triển công nghệ và đổi mới sáng tạo;",
        "Căn cứ Chương trình số 02-CTr/BCĐTW ngày 02 tháng 02 năm 2026 của Ban Chỉ đạo Trung ương về phát triển khoa học, công nghệ, đổi mới sáng tạo và chuyển đổi số về Chương trình công tác năm 2026;",
        "Căn cứ Quyết định số 21/2026/QĐ-TTg ngày 30 tháng 4 năm 2026 của Thủ tướng Chính phủ ban hành Danh mục công nghệ chiến lược và Danh mục sản phẩm công nghệ chiến lược;",
        "Căn cứ Quyết định số 604/QĐ-TTg ngày 02 tháng 4 năm 2026 của Thủ tướng Chính phủ phê duyệt điều chỉnh, bổ sung Chiến lược phát triển khoa học, công nghệ và đổi mới sáng tạo đến năm 2030;",
        "Căn cứ Thông tư số 39/2025/TT-BKHCN ngày 30 tháng 11 năm 2025 của Bộ trưởng Bộ Khoa học và Công nghệ quy định chi tiết và hướng dẫn về lập dự toán, quản lý sử dụng và quyết toán một số nội dung chi ngân sách nhà nước thực hiện nhiệm vụ khoa học, công nghệ và đổi mới sáng tạo;",
        "Căn cứ Thông tư số 44/2025/TT-BKHCN ngày 30 tháng 11 năm 2025 của Bộ trưởng Bộ Khoa học và Công nghệ quy định quản lý nhiệm vụ khoa học và công nghệ do Bộ Khoa học và Công nghệ tài trợ, đặt hàng;",
        "Căn cứ Quyết định số 2227/QĐ-BKHCN ngày 24 tháng 4 năm 2026 của Bộ trưởng Bộ Khoa học và Công nghệ ban hành Kế hoạch tổng thể về khoa học, công nghệ và đổi mới sáng tạo 5 năm giai đoạn 2026-2030;",
        "Căn cứ Quyết định số 2276/QĐ-BKHCN ngày 01 tháng 5 năm 2026 của Bộ trưởng Bộ Khoa học và Công nghệ điều chỉnh nhiệm vụ và phê duyệt kế hoạch tài chính thực hiện tài trợ, đặt hàng, hỗ trợ năm 2026 của Quỹ Phát triển khoa học và công nghệ Quốc gia;",
        "Căn cứ kế hoạch hoạt động năm 2026 và khả năng cân đối kinh phí của Quỹ Phát triển khoa học và công nghệ Quốc gia (sau đây gọi tắt là Quỹ)."
    ]
    
    for cit in citations:
        p_cit = doc.add_paragraph()
        p_cit.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_cit.paragraph_format.space_after = Pt(3)
        p_cit.paragraph_format.line_spacing = 1.15
        run_cit = p_cit.add_run(cit)
        run_cit.italic = True
        run_cit.font.size = Pt(13)
        
    # Spacer
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(6)
    p_space.paragraph_format.space_after = Pt(6)
    p_space.paragraph_format.line_spacing = 1.15
    run_space = p_space.add_run("Quỹ trân trọng thông báo Kế hoạch tài trợ nhiệm vụ nghiên cứu phát triển công nghệ năm 2026, cụ thể như sau:")
    p_space.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # 6. Content Sections
    def add_heading_section(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(13)
        return p
        
    def add_body_paragraph(bold_prefix, text_content):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            run_pre = p.add_run(bold_prefix)
            run_pre.bold = True
            
        p.add_run(text_content)
        return p
        
    def add_bullet_item(text_content):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        
        run_bullet = p.add_run("– ")
        p.add_run(text_content)
        return p

    # Section 1
    add_heading_section("1. Phạm vi và đối tượng nhận tài trợ")
    
    add_body_paragraph("a) Phạm vi tài trợ\n", 
                       "Quỹ xem xét tài trợ các nhiệm vụ nghiên cứu phát triển công nghệ có mục tiêu, nội dung, sản phẩm và khả năng ứng dụng phù hợp với định hướng phát triển khoa học, công nghệ, đổi mới sáng tạo và chuyển đổi số quốc gia.\n\n"
                       "Các nhiệm vụ đăng ký tài trợ tập trung vào một hoặc một số nhóm định hướng sau:")
                       
    tech_bullets = [
        "Công nghệ số, trí tuệ nhân tạo, dữ liệu lớn, bản sao số, điện toán đám mây, điện toán biên, Internet vạn vật, chuỗi khối, an ninh mạng và lượng tử;",
        "Công nghệ mạng di động thế hệ sau, trong đó có 5G, 6G và các công nghệ nền tảng phục vụ hạ tầng kết nối số;",
        "Công nghệ chip bán dẫn, robot, tự động hóa thông minh, thiết bị bay không người lái và các hệ thống tự động;",
        "Công nghệ sinh học, y sinh tiên tiến, thiết bị y tế, dược công nghệ và các giải pháp công nghệ phục vụ chăm sóc sức khỏe;",
        "Công nghệ năng lượng mới, năng lượng tái tạo, lưới điện thông minh, vật liệu mới, vật liệu tiên tiến và công nghệ nguyên tử vì mục đích hòa bình;",
        "Công nghệ biển, đại dương, lòng đất; công nghệ hàng không, vũ trụ; công nghệ đường sắt tốc độ cao, đường sắt đô thị và các công nghệ phục vụ hạ tầng chiến lược;",
        "Công nghệ phục vụ các ngành, lĩnh vực có nhu cầu ứng dụng lớn như điện tử – bán dẫn, logistics, cảng biển, quản lý chuỗi cung ứng thông minh, dịch vụ số cho doanh nghiệp nhỏ và vừa, công nghiệp hỗ trợ, cơ khí chính xác, thủy sản công nghệ cao, nông nghiệp chính xác, hóa chất, vật liệu mới, dệt may và da giày."
    ]
    for tech in tech_bullets:
        add_bullet_item(tech)
        
    add_body_paragraph("b) Loại hình tài trợ: ", "Quỹ xem xét tài trợ nhiệm vụ và cụm nhiệm vụ nghiên cứu phát triển công nghệ. Trong khuôn khổ kêu gọi lần này, Quỹ không tài trợ chuỗi nhiệm vụ.")
    
    add_body_paragraph("c) Đối tượng nhận tài trợ: ", "Tổ chức, doanh nghiệp có tư cách pháp nhân, có năng lực, kinh nghiệm, có chức năng, nhiệm vụ hoặc lĩnh vực hoạt động phù hợp với lĩnh vực nghiên cứu và không thuộc các trường hợp không được xem xét tài trợ thực hiện nhiệm vụ khoa học, công nghệ và đổi mới sáng tạo.")
    
    # Section 2
    add_heading_section("2. Kinh phí tài trợ")
    
    p_budget = add_body_paragraph(None, "Mức trần kinh phí hỗ trợ từ ngân sách nhà nước:")
    add_bullet_item("Nhiệm vụ nghiên cứu phát triển công nghệ tối đa 10 tỷ đồng.")
    add_bullet_item("Đối với cụm nhiệm vụ nghiên cứu phát triển công nghệ tối đa 03 nhiệm vụ/cụm.")
    add_bullet_item("Ngân sách nhà nước hỗ trợ tối đa 50% tổng dự toán kinh phí thực hiện đối với từng nhiệm vụ, từng nhiệm vụ thuộc cụm nhiệm vụ.")
    
    add_body_paragraph(None, "Tổ chức chủ trì phải có phương án, minh chứng huy động nguồn lực hợp pháp ngoài ngân sách nhà nước để bảo đảm phần kinh phí còn lại; Quỹ khuyến khích hồ sơ có tỷ lệ huy động nguồn lực ngoài ngân sách nhà nước cao hơn mức tối thiểu (50%), có nguồn lực cam kết rõ ràng, khả thi và phù hợp với mục tiêu, nội dung, sản phẩm của nhiệm vụ, cụm nhiệm vụ.")
    
    add_body_paragraph(None, "Việc lập dự toán, quản lý, sử dụng và quyết toán kinh phí thực hiện nhiệm vụ, nhiệm vụ thành phần thuộc cụm nhiệm vụ được thực hiện theo Thông tư số 39/2025/TT-BKHCN và các quy định pháp luật có liên quan.")

    # Section 3
    add_heading_section("3. Tiêu chí và yêu cầu đối với nhiệm vụ đăng ký tài trợ")
    
    add_body_paragraph(None, "Nhiệm vụ, cụm nhiệm vụ đăng ký tài trợ phải đáp ứng các tiêu chí quy định tại Điều 6 và khoản 3 Điều 7 Nghị định số 267/2025/NĐ-CP, khoản 3 Điều 4 Thông tư số 44/2025/TT-BKHCN và các quy định có liên quan.")
    add_body_paragraph(None, "Thời gian thực hiện nhiệm vụ, cụm nhiệm vụ tối đa 36 tháng.")
    add_body_paragraph(None, "Nhiệm vụ, cụm nhiệm vụ có thể được xem xét gia hạn 01 lần không quá 12 tháng.")

    # Section 4
    add_heading_section("4. Điều kiện đối với tổ chức chủ trì và nhân lực thực hiện")
    
    add_body_paragraph(None, "Tổ chức đăng ký chủ trì và nhân lực thực hiện nhiệm vụ, cụm nhiệm vụ phải đáp ứng các điều kiện theo quy định của Luật Khoa học, công nghệ và đổi mới sáng tạo, Điều 5 Nghị định số 267/2025/NĐ-CP, khoản 3 Điều 4 Thông tư số 44/2025/TT-BKHCN và các quy định có liên quan.")

    # Section 5
    add_heading_section("5. Hồ sơ đăng ký")
    
    add_body_paragraph(None, "Hồ sơ đăng ký thực hiện nhiệm vụ, cụm nhiệm vụ được lập theo các biểu mẫu quy định tại Điều 11 Nghị định số 267/2025/NĐ-CP, Điều 7 Thông tư số 44/2025/TT-BKHCN, gồm:")
    
    doc_bullets = [
        "Đơn đăng ký chủ trì thực hiện cụm nhiệm vụ (Biểu mẫu BM-02); Đơn đăng ký chủ trì thực hiện nhiệm vụ (Biểu mẫu BM-03);",
        "Thuyết minh nhiệm vụ (Biểu mẫu BM-04);",
        "Lý lịch khoa học của chủ nhiệm nhiệm vụ và các thành viên nghiên cứu (Biểu mẫu BM-06);",
        "Thông tin năng lực và cơ sở vật chất của tổ chức đăng ký chủ trì nhiệm vụ, cụm nhiệm vụ hoặc nhiệm vụ thành phần (Biểu mẫu BM-07);",
        "Văn bản cam kết phối hợp, thử nghiệm, tiếp nhận, ứng dụng, khai thác hoặc thương mại hóa kết quả của doanh nghiệp, cơ quan, tổ chức, địa phương có liên quan, nếu có;",
        "Tài liệu chứng minh phương án huy động vốn đối ứng hoặc nguồn lực ngoài ngân sách nhà nước;",
        "Tài liệu minh chứng quyền sở hữu, quyền sử dụng, quyền khai thác hợp pháp đối với công nghệ, dữ liệu, đối tượng quyền sở hữu trí tuệ; tài liệu về tiêu chuẩn, đo kiểm, kiểm định, thử nghiệm, chứng nhận; tài liệu chứng minh khả năng sử dụng phòng thí nghiệm, cơ sở thử nghiệm, hạ tầng kỹ thuật hoặc các tài liệu chuyên môn khác có liên quan (nếu có);",
        "Các tài liệu khác theo quy định tại Điều 11 Nghị định số 267/2025/NĐ-CP và yêu cầu của biểu mẫu, quy định hiện hành;",
        "Dự toán kinh phí chi tiết thực hiện nhiệm vụ, cụm nhiệm vụ được bổ sung sau khi nhiệm vụ, cụm nhiệm vụ được Hội đồng xét tài trợ đề xuất tài trợ, theo biểu mẫu áp dụng đối với dự toán kinh phí."
    ]
    for db in doc_bullets:
        add_bullet_item(db)

    # Section 6
    add_heading_section("6. Phương thức và thời hạn nộp hồ sơ")
    
    add_body_paragraph("Bước 1: ", "Tổ chức điền hồ sơ trên Hệ thống Quản lý trực tuyến nhiệm vụ khoa học và công nghệ và đổi mới sáng tạo (Hệ thống STM) tại địa chỉ https://stm.mst.gov.vn.")
    add_body_paragraph("Bước 2: ", "Sau khi nộp hồ sơ trên Hệ thống STM, tổ chức lựa chọn một trong hai hình thức nộp hồ sơ sau:")
    
    add_bullet_item("Hình thức 1: Nộp hồ sơ ký số trên Hệ thống STM. Tài liệu ký số tối thiểu gồm Đơn đăng ký, Thuyết minh nhiệm vụ và các văn bản cam kết phối hợp, tiếp nhận, ứng dụng, khai thác hoặc thương mại hóa kết quả, nếu có.")
    add_bullet_item("Hình thức 2: Nộp hồ sơ bản giấy trực tiếp tại Bộ phận Một cửa của Quỹ hoặc gửi qua dịch vụ bưu chính đến Quỹ.")
    
    p_addr = doc.add_paragraph()
    p_addr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_addr.paragraph_format.left_indent = Inches(0.4)
    p_addr.paragraph_format.space_after = Pt(6)
    p_addr.paragraph_format.line_spacing = 1.15
    run_addr_lbl = p_addr.add_run("Địa chỉ tiếp nhận hồ sơ bản giấy: ")
    run_addr_lbl.italic = True
    p_addr.add_run("Bộ phận một cửa Bộ Khoa học và Công nghệ, Tầng 1, Tòa nhà Cục Đổi mới sáng tạo số 113 Trần Duy Hưng, phường Yên Hòa, TP. Hà Nội.\n")
    p_addr.add_run("Trên phong bì ghi rõ: Hồ sơ đề nghị tài trợ nhiệm vụ nghiên cứu phát triển công nghệ năm 2026 và nhóm công nghệ đăng ký tài trợ.")
    
    add_body_paragraph(None, "Trường hợp nộp hồ sơ qua dịch vụ bưu chính, thời điểm nộp hồ sơ được xác định theo dấu bưu chính nơi gửi. Hồ sơ phải được gửi trước thời điểm hết hạn tiếp nhận hồ sơ và gửi đến Quỹ trong thời hạn phù hợp để phục vụ việc rà soát tính hợp lệ theo quy định.")
    add_body_paragraph(None, "Quỹ khuyến khích nộp hồ sơ sử dụng ký số theo chủ trương cải cách thủ tục hành chính, số hóa và minh bạch quản lý nhiệm vụ khoa học, công nghệ và đổi mới sáng tạo.")
    
    p_deadline = doc.add_paragraph()
    p_deadline.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_deadline.paragraph_format.space_after = Pt(6)
    p_deadline.paragraph_format.line_spacing = 1.15
    run_dl_lbl = p_deadline.add_run("Thời hạn tiếp nhận hồ sơ:\n")
    run_dl_lbl.italic = True
    add_bullet_item("Bắt đầu: 8 giờ 00 phút, ngày 10 tháng 7 năm 2026;")
    add_bullet_item("Kết thúc: 17 giờ 00 phút, ngày 10 tháng 8 năm 2026.")

    # Section 7
    add_heading_section("7. Kế hoạch triển khai dự kiến")
    add_bullet_item("Thông báo kết quả xét tài trợ: dự kiến trước ngày 10 tháng 11 năm 2026.")
    
    # Space before sign-off
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # 7. Sign-off block
    sign_table = doc.add_table(rows=1, cols=2)
    sign_table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    sign_table.autofit = False
    
    sign_table.columns[0].width = Inches(3.0)
    sign_table.columns[1].width = Inches(3.4)
    
    cell_recipients = sign_table.cell(0, 0)
    cell_sign = sign_table.cell(0, 1)
    
    # Left cell: Recipients (Nơi nhận)
    p_rec = cell_recipients.paragraphs[0]
    p_rec.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_rec.paragraph_format.line_spacing = 1.0
    p_rec.paragraph_format.space_after = Pt(2)
    
    run_rec_lbl = p_rec.add_run("Nơi nhận:\n")
    run_rec_lbl.bold = True
    run_rec_lbl.italic = True
    run_rec_lbl.font.size = Pt(11)
    
    recipients_list = [
        "- Như Điều 1 (đối tượng);",
        "- HĐQL Quỹ (để b/c);",
        "- Giám đốc Quỹ;",
        "- Lưu: VT, KH."
    ]
    
    for r in recipients_list:
        run_r = p_rec.add_run(r + "\n")
        run_r.font.size = Pt(11)
        
    # Right cell: Signer Role and Name
    p_sign = cell_sign.paragraphs[0]
    p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sign.paragraph_format.line_spacing = 1.1
    p_sign.paragraph_format.space_after = Pt(2)
    
    run_role1 = p_sign.add_run("TM. HỘI ĐỒNG QUẢN LÝ\n")
    run_role1.bold = True
    run_role1.font.size = Pt(12)
    
    run_role2 = p_sign.add_run("CHỦ TỊCH\n\n\n\n\n\n")
    run_role2.bold = True
    run_role2.font.size = Pt(12)
    
    run_name = p_sign.add_run("(Ký tên, đóng dấu)\n\n")
    run_name.italic = True
    run_name.font.size = Pt(11)
    
    # Save the document
    output_path = os.path.join(os.path.dirname(__file__), "Thong_Bao_Ke_Hoach_Tai_Tro_2026.docx")
    doc.save(output_path)
    print(f"Document successfully created at {output_path}")

if __name__ == "__main__":
    create_document()
